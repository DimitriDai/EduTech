# vocab_service/services/grading_service.py
from __future__ import annotations

import os
import re
import json
import time
import uuid
import zipfile
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
import openpyxl

from modules.deepseek_client import DeepSeekClient


@dataclass
class GradeItem:
    no: str
    example_cn: str
    student_en: str
    table_index: int
    row_index: int


@dataclass
class GradeResult:
    no: str
    example_cn: str
    student_en: str

    example: str
    feedback: str

    cache_hit: bool
    ambiguous: bool
    candidates_count: int

    ds_called: bool
    ds_failed: bool
    ds_error: str


# -----------------------------
# Normalization helpers
# -----------------------------
_CN_SPACE_RE = re.compile(r"\s+")
_PUNC_MAP = str.maketrans({
    "，": ",",
    "。": ".",
    "？": "?",
    "！": "!",
    "：": ":",
    "；": ";",
    "（": "(",
    "）": ")",
    "【": "[",
    "】": "]",
    "“": '"',
    "”": '"',
    "‘": "'",
    "’": "'",
})


def normalize_example_cn(s: str) -> str:
    if s is None:
        return ""
    s = str(s).strip().translate(_PUNC_MAP)
    s = _CN_SPACE_RE.sub(" ", s)
    return s.strip()


def normalize_no(s: str) -> str:
    if s is None:
        return ""
    s = str(s).strip()
    # allow "1", "01", "1." etc -> keep digits only if present
    digits = re.findall(r"\d+", s)
    return digits[0] if digits else s


def safe_cell_text(cell) -> str:
    return (cell.text or "").strip()


def detect_header_map(header_cells: List[str]) -> Optional[Dict[str, int]]:
    headers = [c.strip() for c in header_cells]

    def find_idx(keys: List[str]) -> Optional[int]:
        for i, h in enumerate(headers):
            for k in keys:
                if k in h:
                    return i
        return None

    example_cn_i = find_idx(["例句翻译", "中文翻译", "翻译", "中文"])
    if example_cn_i is None:
        return None

    no_i = find_idx(["No", "NO", "序号"])
    if no_i is None:
        no_i = 0

    student_i = find_idx(["例句", "学生", "作答", "答案"])
    if student_i is None:
        student_i = 2 if len(headers) >= 3 else len(headers) - 1

    return {"no": no_i, "example_cn": example_cn_i, "student_en": student_i}


def extract_items_from_docx(docx_path: str) -> Tuple[List[GradeItem], Dict[str, Any]]:
    doc = Document(docx_path)
    items: List[GradeItem] = []

    meta = {
        "docx_path": docx_path,
        "tables_total": len(doc.tables),
        "tables_parsed": 0,
        "rows_total": 0,
        "rows_candidate": 0,
        "rows_skipped": 0,
        "skipped_reasons": {},
    }

    def skip(reason: str):
        meta["rows_skipped"] += 1
        meta["skipped_reasons"][reason] = meta["skipped_reasons"].get(reason, 0) + 1

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        meta["rows_total"] += len(table.rows)

        header_map = None
        header_row_idx = None
        scan_limit = min(3, len(table.rows))
        for r in range(scan_limit):
            headers = [safe_cell_text(c) for c in table.rows[r].cells]
            m = detect_header_map(headers)
            if m:
                header_map = m
                header_row_idx = r
                break

        if not header_map:
            continue

        meta["tables_parsed"] += 1

        for r_idx in range(header_row_idx + 1, len(table.rows)):
            row = table.rows[r_idx]
            cells = [safe_cell_text(c) for c in row.cells]
            meta["rows_candidate"] += 1

            try:
                no = cells[header_map["no"]] if header_map["no"] < len(cells) else ""
                example_cn = cells[header_map["example_cn"]] if header_map["example_cn"] < len(cells) else ""
                student_en = cells[header_map["student_en"]] if header_map["student_en"] < len(cells) else ""
            except Exception:
                skip("row_parse_error")
                continue

            if not normalize_example_cn(example_cn):
                skip("example_cn_empty")
                continue

            items.append(
                GradeItem(
                    no=str(no).strip(),
                    example_cn=str(example_cn).strip(),
                    student_en=str(student_en).strip(),
                    table_index=t_idx,
                    row_index=r_idx,
                )
            )

    return items, meta


# -----------------------------
# Index building (FAST path: shuffle master xlsx)
# -----------------------------

# Index value list: (example_en, source)
ExampleIndex = Dict[str, List[Tuple[str, str]]]
ExampleIndexByNo = Dict[Tuple[str, str], List[Tuple[str, str]]]


def build_example_index_from_shuffle_master_xlsx(xlsx_path: str) -> Tuple[ExampleIndex, ExampleIndexByNo, Dict[str, Any]]:
    """
    读取 shuffle_e2c_master.xlsx（多个 sheet）：
    - key: normalize(example_cn) -> example
    - key2: (no, normalize(example_cn)) -> example
    返回： (idx_cn, idx_no_cn, meta)
    """
    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"shuffle master xlsx not found: {xlsx_path}")

    wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)

    idx_cn: ExampleIndex = {}
    idx_no_cn: ExampleIndexByNo = {}

    meta = {
        "xlsx_path": xlsx_path,
        "sheets": wb.sheetnames,
        "rows_read": 0,
        "rows_used": 0,
        "rows_skipped": 0,
        "skipped_reasons": {},
    }

    def skip(reason: str):
        meta["rows_skipped"] += 1
        meta["skipped_reasons"][reason] = meta["skipped_reasons"].get(reason, 0) + 1

    for sh in wb.sheetnames:
        ws = wb[sh]
        # read header row
        headers = []
        for c in range(1, ws.max_column + 1):
            v = ws.cell(1, c).value
            if v is None:
                headers.append("")
            else:
                headers.append(str(v).strip())

        # required columns
        def col_i(name: str) -> Optional[int]:
            try:
                return headers.index(name) + 1
            except ValueError:
                return None

        c_example = col_i("example")
        c_example_cn = col_i("example_cn")
        c_no = col_i("no")

        if c_example is None or c_example_cn is None:
            # not a valid shuffle sheet
            continue

        for r in range(2, ws.max_row + 1):
            meta["rows_read"] += 1

            example = ws.cell(r, c_example).value
            example_cn = ws.cell(r, c_example_cn).value
            no = ws.cell(r, c_no).value if c_no is not None else ""

            example = "" if example is None else str(example).strip()
            example_cn_raw = "" if example_cn is None else str(example_cn).strip()

            k_cn = normalize_example_cn(example_cn_raw)
            if not k_cn:
                skip("example_cn_empty")
                continue
            if not example:
                skip("example_empty")
                continue

            no_norm = normalize_no(no)

            idx_cn.setdefault(k_cn, []).append((example, "shuffle_master"))
            idx_no_cn.setdefault((no_norm, k_cn), []).append((example, "shuffle_master"))

            meta["rows_used"] += 1

    return idx_cn, idx_no_cn, meta


# -----------------------------
# Index building (fallback: cache json)
# -----------------------------

def load_json(path: str) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _iter_entry_dicts(obj: Any) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []

    def walk(x: Any):
        if isinstance(x, dict):
            has_cn = ("example_cn" in x) or ("exampleCn" in x)
            has_en = ("example" in x) or ("example_en" in x)
            if has_cn and has_en:
                out.append(x)
            for v in x.values():
                walk(v)
        elif isinstance(x, list):
            for v in x:
                walk(v)

    walk(obj)
    return out


def build_example_index_from_cache(
    global_cache_path: str,
    uploaded_cache_path: Optional[str] = None,
) -> ExampleIndex:
    index: ExampleIndex = {}

    if os.path.exists(global_cache_path):
        g = load_json(global_cache_path)
        for e in _iter_entry_dicts(g):
            ex_cn = e.get("example_cn") or e.get("exampleCn") or ""
            ex_en = e.get("example") or e.get("example_en") or ""
            k = normalize_example_cn(ex_cn)
            if k and str(ex_en).strip():
                index.setdefault(k, []).append((str(ex_en).strip(), "global"))

    if uploaded_cache_path and os.path.exists(uploaded_cache_path):
        u = load_json(uploaded_cache_path)
        for e in _iter_entry_dicts(u):
            ex_cn = e.get("example_cn") or e.get("exampleCn") or ""
            ex_en = e.get("example") or e.get("example_en") or ""
            k = normalize_example_cn(ex_cn)
            if k and str(ex_en).strip():
                index.setdefault(k, []).append((str(ex_en).strip(), "uploaded"))

    return index


def pick_best_example(candidates: List[Tuple[str, str]]) -> Tuple[str, bool, int]:
    if not candidates:
        return "", False, 0

    count = len(candidates)

    # prefer global, then shuffle_master, then others
    def source_rank(src: str) -> int:
        if src == "global":
            return 0
        if src == "shuffle_master":
            return 1
        if src == "uploaded":
            return 2
        return 9

    candidates_sorted = sorted(
        candidates,
        key=lambda x: (source_rank(x[1]), -len(x[0] or "")),
    )
    best = (candidates_sorted[0][0] or "").strip()
    ambiguous = count > 1
    return best, ambiguous, count


# -----------------------------
# Prompt & grading
# -----------------------------

def build_prompt(example_cn: str, student_en: str, ref_en: str) -> str:
    ref_part = ""
    if (ref_en or "").strip():
        ref_part = f"\n\n【参考英文】\n{ref_en.strip()}\n"

    return f"""
你是一名英文改错老师，请根据给定中文和学生英文完成改错。

【中文】
{example_cn}

【学生英文】
{student_en}{ref_part}

【要求】
表达中文的关键信息，但是没必要字字对应，重要的是语法正确性！
不用非得匹配标注的参考例句，保证语法语义正确即可！
指出所有语法、用词、拼写、时态、语序等问题
给出的正确版本应在尽量保留学生句式的前提下进行“最小修改”

【输出格式（务必严格遵守，不要添加多余内容）】
【错误】这里写所有错误（若无错误请写：无错误）
【正确版本】这里写修改后的英文句子
""".strip()


def grade_docx(
    docx_path: str,
    example_index: ExampleIndex,
    example_index_by_no: Optional[ExampleIndexByNo] = None,
    *,
    request_id: Optional[str] = None,
    timeout_s: int = 30,
    retries: int = 2,
    retry_backoff_sec: float = 1.0,
    temperature: float = 0.2,
    max_tokens: int = 320,
    use_ref_when_blank: bool = True,
) -> Tuple[List[GradeResult], Dict[str, Any]]:
    request_id = request_id or str(uuid.uuid4())

    client = DeepSeekClient(timeout=timeout_s, retries=retries, retry_backoff_sec=retry_backoff_sec)
    items, parse_meta = extract_items_from_docx(docx_path)

    meta: Dict[str, Any] = {
        "request_id": request_id,
        "docx_path": docx_path,
        "items_total": len(items),

        "graded_count": 0,
        "skipped_blank_student": 0,

        "hit_by_no_cn": 0,
        "hit_by_cn_only": 0,
        "cache_miss": 0,

        "ambiguous": 0,

        "ds_called": 0,
        "ds_failed": 0,

        "parse": parse_meta,
    }

    results: List[GradeResult] = []

    for it in items:
        k_cn = normalize_example_cn(it.example_cn)
        no_norm = normalize_no(it.no)

        candidates: List[Tuple[str, str]] = []
        hit_mode = "miss"

        if example_index_by_no is not None:
            candidates = example_index_by_no.get((no_norm, k_cn), [])
            if candidates:
                hit_mode = "no_cn"

        if not candidates:
            candidates = example_index.get(k_cn, [])
            if candidates:
                hit_mode = "cn_only"

        example, ambiguous, cand_count = pick_best_example(candidates)

        if hit_mode == "no_cn":
            meta["hit_by_no_cn"] += 1
        elif hit_mode == "cn_only":
            meta["hit_by_cn_only"] += 1
        else:
            meta["cache_miss"] += 1

        if ambiguous:
            meta["ambiguous"] += 1

        if not (it.student_en or "").strip():
            meta["skipped_blank_student"] += 1
            feedback = "【错误】未作答\n【正确版本】"
            if use_ref_when_blank and (example or "").strip():
                feedback = f"【错误】未作答\n【正确版本】{example.strip()}"
            results.append(
                GradeResult(
                    no=it.no,
                    example_cn=it.example_cn,
                    student_en=it.student_en,
                    example=example,
                    feedback=feedback,
                    cache_hit=bool(example),
                    ambiguous=ambiguous,
                    candidates_count=cand_count,
                    ds_called=False,
                    ds_failed=False,
                    ds_error="",
                )
            )
            continue

        prompt = build_prompt(it.example_cn, it.student_en, example)
        meta["ds_called"] += 1

        try:
            feedback = client.call_model(prompt=prompt, max_tokens=max_tokens, temperature=temperature)
            meta["graded_count"] += 1
            results.append(
                GradeResult(
                    no=it.no,
                    example_cn=it.example_cn,
                    student_en=it.student_en,
                    example=example,
                    feedback=(feedback or "").strip(),
                    cache_hit=bool(example),
                    ambiguous=ambiguous,
                    candidates_count=cand_count,
                    ds_called=True,
                    ds_failed=False,
                    ds_error="",
                )
            )
        except Exception as e:
            meta["ds_failed"] += 1
            results.append(
                GradeResult(
                    no=it.no,
                    example_cn=it.example_cn,
                    student_en=it.student_en,
                    example=example,
                    feedback="【错误】批改失败（系统原因）\n【正确版本】",
                    cache_hit=bool(example),
                    ambiguous=ambiguous,
                    candidates_count=cand_count,
                    ds_called=True,
                    ds_failed=True,
                    ds_error=str(e)[:300],
                )
            )

        time.sleep(0.02)  # 更轻的节流（需要的话你也可以直接删掉）

    return results, meta


# -----------------------------
# Batch helpers (zip)
# -----------------------------

def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def is_zip(path: str) -> bool:
    return path.lower().endswith(".zip")


def extract_zip(zip_path: str, to_dir: str) -> List[str]:
    ensure_dir(to_dir)
    with zipfile.ZipFile(zip_path, "r") as z:
        z.extractall(to_dir)

    docxs: List[str] = []
    for root, _, files in os.walk(to_dir):
        for fn in files:
            if fn.lower().endswith(".docx") and not fn.startswith("~$"):
                docxs.append(os.path.join(root, fn))
    return sorted(docxs)


def make_zip(zip_path: str, files: List[str]) -> None:
    ensure_dir(os.path.dirname(zip_path) or ".")
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, arcname=os.path.basename(f))
