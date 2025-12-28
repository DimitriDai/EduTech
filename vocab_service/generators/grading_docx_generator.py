# vocab_service/generators/grading_docx_generator.py
from __future__ import annotations

import os
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.grading_service import (
    GradeResult,
    normalize_example_cn,
    safe_cell_text,
    detect_header_map,
)

# ===================== 显性展示数值需求（按你给的定死） =====================

# 列宽（cm）：序号, 中文, 中译英, 批改, 正确例句
COL_WIDTHS = [1.5, 4, 4, 10, 4]

# 行高（cm）
HEADER_HEIGHT = 1.0
DATA_HEIGHT = 4.0

FONT_NAME = "等线"
FONT_SIZE = 11


# ===================== 公共工具函数（按你的风格） =====================

def clear_cell(cell):
    """清空单元格全部内容（包括默认空段落）。"""
    cell._element.clear_content()


def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(height_cm * 567)))  # 1cm ≈ 567 twips
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def make_landscape(document: Document):
    """将文档设为横版。"""
    for section in document.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        w, h = section.page_height, section.page_width
        section.page_width = w
        section.page_height = h


def add_page_number(document: Document):
    """页脚添加居中页码。"""
    for section in document.sections:
        footer = section.footer
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def set_cell_font(cell):
    """统一设置单元格内字体为 等线 11。"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            # east asia font
            if run._element.rPr is None:
                rPr = OxmlElement("w:rPr")
                run._element.insert(0, rPr)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_table_col_widths_cm(table, widths_cm: List[float]) -> None:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(float(widths_cm[i]))


def _clean_text(s: str) -> str:
    """
    去除多余空行：把连续空白压成单个换行；
    同时 strip。
    """
    if s is None:
        return ""
    s = str(s)
    s = re.sub(r"\r\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)  # 最多保留双换行
    return s.strip()


def _write_cell_text(cell, text: str) -> None:
    """
    先 clear_cell，再写入单段落/多段落（按换行拆）。
    目标：不保留 docx 默认空段落，且尽量不出现“空行”。
    """
    clear_cell(cell)
    text = _clean_text(text)

    if not text:
        # 保持空
        return

    lines = text.split("\n")
    # 第一段落
    p0 = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p0.text = lines[0]

    # 后续段落（不写空段）
    for line in lines[1:]:
        if line.strip() == "":
            continue
        cell.add_paragraph(line)

    set_cell_font(cell)


def _results_map(results: List[GradeResult]) -> Dict[str, GradeResult]:
    """
    以 normalize(example_cn) 为键。
    同一句中文重复时保留首次出现（与原练习表更一致）。
    """
    m: Dict[str, GradeResult] = {}
    for r in results:
        k = normalize_example_cn(r.example_cn)
        if k and k not in m:
            m[k] = r
    return m


def write_feedback_docx(
    input_docx_path: str,
    output_docx_path: str,
    results: List[GradeResult],
) -> None:
    """
    生成“批改反馈 docx”（严格按 review_ch_to_en.py 的展示数值规范）：
    - 横版
    - 页码
    - 等线 11
    - 列宽：1.5 / 4 / 4 / 10 / 4 cm
    - 行高 exact：表头 1.0cm，数据 3.8cm
    - 清空单元格默认空段落，尽量去掉空行
    - 每个识别到的练习表格，在输出文档中生成对应的新表（不修改原表）
      表头：序号 | 中文 | 中译英 | 批改 | 正确例句
    """
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)

    src = Document(input_docx_path)
    out = Document()

    # 横版 + 页码
    make_landscape(out)
    add_page_number(out)

    res_map = _results_map(results)

    tables_created = 0

    for table in src.tables:
        if not table.rows:
            continue

        header_map = None
        header_row_idx = None

        scan_limit = min(3, len(table.rows))
        for rr in range(scan_limit):
            headers = [safe_cell_text(c) for c in table.rows[rr].cells]
            m = detect_header_map(headers)
            if m:
                header_map = m
                header_row_idx = rr
                break

        if not header_map:
            continue

        # 统计有效数据行（example_cn 非空）
        data_rows: List[Tuple[str, str, str]] = []  # (no, example_cn, student_en)
        for r_idx in range(header_row_idx + 1, len(table.rows)):
            row = table.rows[r_idx]
            cells = [safe_cell_text(c) for c in row.cells]

            no = cells[header_map["no"]] if header_map["no"] < len(cells) else ""
            example_cn = cells[header_map["example_cn"]] if header_map["example_cn"] < len(cells) else ""
            student_en = cells[header_map["student_en"]] if header_map["student_en"] < len(cells) else ""

            if not normalize_example_cn(example_cn):
                continue

            data_rows.append((str(no).strip(), str(example_cn).strip(), str(student_en).strip()))

        if not data_rows:
            continue

        # 输出文档：每个源表格生成一个新表
        if tables_created > 0:
            out.add_paragraph("")  # 表格间留一点点间距（不影响单元格空行规则）
        
        # 创建表格
        out_table = out.add_table(rows=1 + len(data_rows), cols=5)
        out_table.style = "Table Grid"
        _set_table_col_widths_cm(out_table, COL_WIDTHS)

        # 写表头
        hdr = out_table.rows[0]
        set_row_height(hdr, HEADER_HEIGHT)

        headers = ["序号", "中文", "中译英", "批改", "正确例句"]
        for i, htxt in enumerate(headers):
            _write_cell_text(hdr.cells[i], htxt)

        # 写数据行
        for i, (no, example_cn, student_en) in enumerate(data_rows, start=1):
            row = out_table.rows[i]
            set_row_height(row, DATA_HEIGHT)

            k = normalize_example_cn(example_cn)
            gr = res_map.get(k)

            # 你的列含义要求：
            # 序号 | 中文(example_cn) | 中译英(student_en) | 批改(feedback/错误讲解) | 正确例句(example)
            _write_cell_text(row.cells[0], no)
            _write_cell_text(row.cells[1], example_cn)
            _write_cell_text(row.cells[2], student_en)

            if gr:
                # 批改列：写 DS 原文（包含【错误】【正确版本】）
                _write_cell_text(row.cells[3], gr.feedback or "")
                # 正确例句列：写 example（永远来自答案库/缓存匹配）
                _write_cell_text(row.cells[4], gr.example or "")
            else:
                _write_cell_text(row.cells[3], "【错误】未匹配到批改结果\n【正确版本】")
                _write_cell_text(row.cells[4], "")

        tables_created += 1

    if tables_created == 0:
        # 如果源 docx 没识别到任何练习表，输出一个提示（也按字体写）
        p = out.add_paragraph("未识别到练习表格。")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    out.save(output_docx_path)