# ============================================================
# formatter_ai.py  （完整覆盖版）
#
# 你可以直接把本文件整体复制覆盖到：
# C:\Users\24340\Desktop\EduTech\streamlit_wrt_app\Wrt_requirements\modules\formatter_ai.py
#
# 功能：
# 1) DeepSeek 生成“逐句批改反馈”（文本）
# 2) DeepSeek 生成“AI 主裁评分”（严格 JSON）
#    - 单项 TR/CC/LR/GRA 只能整数
#    - overall = 四项平均后 向下取到最近 0.5（6.25->6.0; 6.75->6.5）
# 3) 生成 Word 报告（docx）：
#    - 原文 vs 优化后对比表格
#    - 逐句反馈
#    - 段落优化（如模型输出了）
#    - 评分建议（可选）
#    - 【改错练习】（自动生成）
#
# 稳态增强：
# - timeout + 重试 + 失败降级（避免 500）
# - JSON 提取容错（避免模型返回多余文本导致解析失败）
# ============================================================

import os
import re
import json
import time
from typing import Optional, Dict, Any, Literal

import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ============================================================
# 0) DeepSeek 配置
# ============================================================

DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"

def _get_deepseek_key(api_key: Optional[str] = None) -> str:
    # 优先用传入参数，其次环境变量
    key = (api_key or os.environ.get("DEEPSEEK_API_KEY", "") or "").strip()
    return key


def _post_with_retry(url: str, headers: dict, payload: dict, tries: int = 3, timeout: int = 120) -> requests.Response:
    last_err = None
    for attempt in range(tries):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
            resp.raise_for_status()
            return resp
        except Exception as e:
            last_err = e
            time.sleep(1 * (2 ** attempt))  # 1s,2s,4s
    raise last_err


def _extract_json_object(text: str) -> str:
    """
    从模型输出中提取第一个 { ... } JSON 对象字符串
    允许有 ```json 包裹或少量多余文本。
    """
    s = (text or "").strip()
    s = re.sub(r"^```json\s*", "", s, flags=re.I).strip()
    s = re.sub(r"^```\s*", "", s).strip()
    s = re.sub(r"\s*```$", "", s).strip()

    start = s.find("{")
    end = s.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("AI 返回内容中找不到 JSON 对象")
    return s[start:end + 1]


# ============================================================
# 1) 基础工具：读文件 / 猜 Task
# ============================================================

def load_text_from_file(filepath: str) -> str:
    """
    只支持 .txt / .docx
    """
    fp = filepath.lower()
    if fp.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    if fp.endswith(".docx"):
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    raise ValueError("仅支持 .txt 或 .docx 文件（.doc 暂不支持）")


def guess_task_type(text: str) -> Literal["Task 1", "Task 2"]:
    """
    简单启发式猜测，小作文通常包含图表类关键词
    """
    t = (text or "").lower()
    task1_keywords = ["chart", "table", "graph", "diagram", "map"]
    if any(k in t for k in task1_keywords):
        return "Task 1"
    return "Task 2"

def detect_task_type_via_deepseek(text: str, api_key: Optional[str] = None) -> str:
    """
    用 DeepSeek 做二分类：只返回 'Task 1' 或 'Task 2'
    安全策略：
    - temperature=0，max_tokens 极小
    - 输出不合法/失败 => 返回 'Task 1'（不补写，最安全）
    """
    key = _get_deepseek_key(api_key)
    if not key:
        return "Task 1"

    t = (text or "").strip()
    if not t:
        return "Task 1"

    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}

    prompt = f"""
You are an IELTS examiner. Classify the following essay as either Task 1 or Task 2.

Rules:
- Output MUST be exactly one of: Task 1 / Task 2
- No punctuation, no explanation, no extra words.

Essay:
{t}
""".strip()

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.0,
        "max_tokens": 8,
    }

    try:
        r = _post_with_retry(DEEPSEEK_URL, headers, payload, tries=2, timeout=30)
        out = (r.json()["choices"][0]["message"]["content"] or "").strip()
        out = out.replace("\u200b", "").strip()

        if out == "Task 1" or out == "Task 2":
            return out

        # 容错：有些模型会返回 "Task 2\n" 或多一行
        first_line = out.splitlines()[0].strip() if out else ""
        if first_line in ("Task 1", "Task 2"):
            return first_line

        return "Task 1"
    except Exception:
        return "Task 1"

# ============================================================
# 2) DeepSeek：逐句批改反馈（文本）
# ============================================================

def get_feedback_from_deepseek(text: str, api_key: Optional[str] = None, task_type: str = "Task 2") -> str:
    """
    输出：纯文本反馈（不要求 JSON）
    失败：返回可读的降级提示（不抛异常，避免后端 500）
    """
    key = _get_deepseek_key(api_key)
    if not key:
        return "[AI反馈生成失败：缺少 DEEPSEEK_API_KEY]"

    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}

    prompt = f"""
你是 IELTS 写作老师。请对学生作文进行逐句反馈与修正，并在文末给出“段落优化”和“优化依据”。

总原则（必须严格遵守）：
A) i+1 原则：每个句子的【修正】只能做“一个维度”的升级：
   - 维度1：词汇/搭配升级（更准确或更地道，但不能明显更学术、更生僻）
   - 维度2：语法/句式升级（只做一个点：如从简单句→含从句；或更自然的连接；但不能连着升两三层）
   每句只允许选其一，绝不允许同时做词汇升级 + 句式升级。
B) 难度上限：所有【修正】必须与学生原句难度“相近略高”，禁止显著拔高（禁止突然出现复杂从句链、学术词、长难句堆叠）。
C) 保留学生表达意图：不得改变原句核心意思、立场、信息点；只修错与轻度提升。
D) 如果一句话本来没问题：必须写【错误】=无，【修正】=无（不要为了“显得更好”强行改）。

输出格式要求（必须严格按下面结构，不要额外内容）：
1) 每个原句前加编号：1. 2. 3. ...
2) 每句后必须有两行：
   【错误】：“类型 + 具体点”的方式说明（如无写“无”）。
           类型可选：语法/时态/主谓一致/冠词/介词/拼写/搭配/用词不准确/句子结构/衔接与标点/重复与冗余/逻辑不清
           必须指出具体位置或关键词（例如：damage to 多余；With using 结构错误；plating 拼写错误）。
   【修正】：xxx（如【错误】为“无”，【修正】写“无”；并严格遵守 i+1 规则）
3) 全部句子结束后，输出：
   【段落优化】：输出一篇“可直接交卷”的完整优化稿，但【难度必须与逐句修正保持同一档，不得比逐句修正更高级】；如无需整体改写，输出“无”。
   【优化依据】：用要点列出你做过的修改类型（例如：时态一致、冠词、搭配、连接词、避免重复等），不要长篇解释。
4) 仅输出上述内容，不要寒暄。

作文类型：{task_type}
学生作文：
{text.strip()}
""".strip()

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
    }

    try:
        r = _post_with_retry(DEEPSEEK_URL, headers, payload, tries=3, timeout=120)
        return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"[AI反馈生成失败：{type(e).__name__}] {str(e)[:200]}"


# ============================================================
# 3) DeepSeek：AI 主裁评分（严格 JSON）
#    - 单项整数
#    - overall = avg 向下取 0.5（写作单项规则）
# ============================================================

def get_ai_band_scores_from_deepseek(text: str, task_type: str, api_key: Optional[str] = None) -> Dict[str, Any]:
    """
    返回 dict；失败返回 {}（不抛异常，保证稳态）
    """
    key = _get_deepseek_key(api_key)
    if not key:
        return {}

    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}

    prompt = f"""
你是一位符合真实雅思官方评分尺度的 IELTS Writing Examiner。
请对下面作文给出四项分数与总分。输出必须是严格 JSON（不要 markdown，不要任何多余文字）。

必须严格遵守的规则：
1) 四个单项（TR/TA、CC、LR、GRA）只能是整数：4,5,6,7,8,9（不允许 0.5）
2) Writing overall 的计算方式：
   - avg = (TR + CC + LR + GRA) / 4
   - overall = 向下取到最近的 0.5 档（floor to 0.5）
     例：6.25 -> 6.0；6.75 -> 6.5；7.00 -> 7.0
3) 评分方法：以“主体特征”判断档位，不要用“错误计数法”随意压一档。
4) reasons：每项给 2~4 条中文要点，简短可操作；overall reasons 1~2 条即可。

输出 JSON（字段名必须完全一致）：
{{
  "task_type": "{task_type}",
  "TR": {{"score": 7, "reasons": ["...","..."]}},
  "CC": {{"score": 7, "reasons": ["...","..."]}},
  "LR": {{"score": 6, "reasons": ["...","..."]}},
  "GRA": {{"score": 6, "reasons": ["...","..."]}},
  "overall": {{"score": 6.5, "avg": 6.75, "reasons": ["..."]}}
}}

作文类型：{task_type}
学生作文：
{text.strip()}
""".strip()

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
    }

    try:
        r = _post_with_retry(DEEPSEEK_URL, headers, payload, tries=3, timeout=120)
        raw = r.json()["choices"][0]["message"]["content"]
        json_str = _extract_json_object(raw)
        data = json.loads(json_str)
    except Exception:
        return {}

    # ---- 评分兜底校验：保证输出永远合法（单项整数 + overall 重算） ----
    def _safe_int_score(v) -> Optional[int]:
        try:
            x = float(v)
            # 单项必须整数：这里“就地纠偏”——直接取 floor
            # （不要在这里做复杂抬分，否则你会更难控制）
            return int(x)
        except Exception:
            return None

    tr = _safe_int_score(data.get("TR", {}).get("score"))
    cc = _safe_int_score(data.get("CC", {}).get("score"))
    lr = _safe_int_score(data.get("LR", {}).get("score"))
    gra = _safe_int_score(data.get("GRA", {}).get("score"))

    # 如果四项齐全，重算 overall（向下取 0.5）
    scores = [tr, cc, lr, gra]
    if all(s is not None for s in scores):
        avg = sum(scores) / 4.0
        overall = int(avg * 2) / 2.0  # floor to 0.5
        data.setdefault("overall", {})
        data["overall"]["avg"] = avg
        data["overall"]["score"] = overall

        # 也把单项回填为整数（确保输出合规）
        data.setdefault("TR", {}).update({"score": tr})
        data.setdefault("CC", {}).update({"score": cc})
        data.setdefault("LR", {}).update({"score": lr})
        data.setdefault("GRA", {}).update({"score": gra})

    return data


# ============================================================
# 4) 解析反馈：抽取“修正句子/段落优化/改错练习”
# ============================================================

def _parse_feedback_blocks(feedback_text: str) -> Dict[str, Any]:
    """
    从反馈文本中解析：
    - items: [{"no":1, "original":"...", "error":"...", "fix":"..."}...]
    - optimized: 段落优化（可能是"无"）
    - rationale: 优化依据（可能为空）
    """
    text = feedback_text or ""
    lines = [ln.rstrip() for ln in text.splitlines()]

    # 解析编号句子块
    items = []
    cur_no = None
    cur_original = None
    cur_error = None
    cur_fix = None

    def flush():
        nonlocal cur_no, cur_original, cur_error, cur_fix
        if cur_no is not None and cur_original is not None:
            items.append({
                "no": cur_no,
                "original": (cur_original or "").strip(),
                "error": (cur_error or "").strip() if cur_error else "",
                "fix": (cur_fix or "").strip() if cur_fix else "",
            })
        cur_no = None
        cur_original = None
        cur_error = None
        cur_fix = None

    # 支持 "1." "1、" "1)" 等
    num_pat = re.compile(r"^\s*(\d+)\s*[\.、\)]\s*(.*)$")

    optimized = ""
    rationale = ""
    in_optimized = False
    in_rationale = False

    for ln in lines:
        # 进入“段落优化”块（兼容同行带内容）
        if "【段落优化】" in ln:
            in_optimized = True
            in_rationale = False
            after = ln.split("】", 1)[-1].lstrip("：:").strip()
            if after:
                optimized += after + "\n"
            continue

        # 进入“优化依据”块（兼容同行带内容）
        if "【优化依据】" in ln:
            in_rationale = True
            in_optimized = False
            after = ln.split("】", 1)[-1].lstrip("：:").strip()
            if after:
                rationale += after + "\n"
            continue

        # 正在收集段落优化 / 优化依据
        if in_optimized:
            optimized += (ln + "\n")
            continue
        if in_rationale:
            rationale += (ln + "\n")
            continue

        # 解析 numbered sentence
        m = num_pat.match(ln)
        if m:
            flush()
            cur_no = int(m.group(1))
            cur_original = m.group(2).strip()
            continue

        if "【错误】" in ln:
            cur_error = ln.split("】", 1)[-1].strip()
            cur_error = re.sub(r"^[：:\-\–\—•\u2022]+\s*", "", cur_error)
            continue

        if "【修正】" in ln:
            cur_fix = ln.split("】", 1)[-1].strip()
            # 清洗：去掉模型偶尔吐出的行首符号（全角/半角冒号、项目符号、破折号等）
            cur_fix = re.sub(r"^[：:\-\–\—•\u2022]+\s*", "", cur_fix)
            continue

    flush()

    optimized = optimized.strip()
    rationale = rationale.strip()

    if optimized in ["无", "无。", "none", "None"]:
        optimized = ""

    return {
        "items": items,
        "optimized": optimized,
        "rationale": rationale,
    }


def _build_optimized_text_from_items(items: list) -> str:
    """
    用编号句子的【修正】拼接出“优化后文章”。
    规则：
    - 若【修正】为“无”，则使用 original
    """
    fixed_lines = []
    for it in items:
        orig = (it.get("original") or "").strip()
        fix = (it.get("fix") or "").strip()
        if fix in ["无", "无。", "", "none", "None"]:
            fixed_lines.append(orig)
        else:
            fixed_lines.append(fix)
    # 用换行拼接更像作文段落（模型通常按句输出，不会带段落；这里保持可读）
    return "\n".join([x for x in fixed_lines if x.strip()])


def _extract_error_drills(items: list) -> list:
    """
    从 items 里提取“改错练习”条目：
    - 只要【错误】不是“无”，就收录
    输出：[{no, original, error}]
    """
    drills = []
    for it in items:
        err = (it.get("error") or "").strip()
        if not err or err in ["无", "无。", "none", "None"]:
            continue
        drills.append({
            "no": it.get("no"),
            "original": (it.get("original") or "").strip(),
            "error": err
        })
    return drills

def count_words_rough(text: str) -> int:
    """
    粗略英文词数统计：按字母/数字序列计数，稳定、够用。
    """
    if not text:
        return 0
    tokens = re.findall(r"[A-Za-z0-9]+(?:'[A-Za-z0-9]+)?", text)
    return len(tokens)

def generate_optional_example_with_deepseek(
    essay_text: str,
    task_type: str,
    api_key: Optional[str] = None,
    max_tokens: int = 160,
) -> str:
    """
    只生成“可选补充例子/细节”本身（1个例子 + 1-2句解释），不重写全文。
    失败返回空字符串。
    """
    base = (essay_text or "").strip()
    if not base:
        return ""

    key = _get_deepseek_key(api_key)
    if not key:
        return ""

    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}

    prompt = f"""
你是一位 IELTS 写作老师。下面是一篇已优化过的 Task 2 作文。
请只输出一个“可选补充例子/细节”（用于补足字数或增强说服力），要求：

1) 只输出新增内容本身：1 段（或 2–4 句），不要重写全文、不要复述原文、不要输出标题。
2) 难度锁定：用与原文相同档次的词汇与句式，不得升级到更学术/更复杂。
3) 内容必须与原文论点一致，可直接插入到正文某一段作为例子/解释。
4) 禁止输出任何格式标记（不要编号、不要“【】”、不要解释你在做什么）。

作文类型：{task_type}
原文：
{base}
""".strip()

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
        "max_tokens": max_tokens,
    }

    try:
        r = _post_with_retry(DEEPSEEK_URL, headers, payload, tries=2, timeout=60)
        out = (r.json()["choices"][0]["message"]["content"] or "").strip()
        out = strip_markdown_marks(out).strip()

        # 简单清洗：如果模型没按要求，尽量兜底
        # 若返回过长（像在重写全文），直接丢弃
        if count_words_rough(out) > 160:
            return ""

        # 若返回太短（<8词）也没意义
        if count_words_rough(out) < 12:
            return ""

        return out
    except Exception:
        return ""

def _looks_english_sentence(s: str) -> bool:
    """
    用于“是否像英文句子”的稳健判定：
    - 至少含 3 个英文字母
    - 汉字占比不能太高
    目的：遇到中文/乱码/混合，直接回退 original，避免污染优化稿。
    """
    if not s:
        return False
    s = s.strip()
    letters = len(re.findall(r"[A-Za-z]", s))
    if letters < 3:
        return False
    cjk = len(re.findall(r"[\u4e00-\u9fff]", s))
    # 允许极少量中文（比如括号里备注），但不能太多
    return cjk <= 2

# ============================================================
# 5) Word 输出：对比表格 + 反馈 + 段落优化 + 改错练习
# ============================================================

def _set_doc_default_font(doc: Document, font_name: str = "等线", size_pt: int = 11):
    from docx.oxml.ns import qn

    def apply_font(style):
        style.font.name = font_name
        style.font.size = Pt(size_pt)
        rFonts = style._element.rPr.rFonts
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)

    # Normal
    apply_font(doc.styles["Normal"])

    # 所有 Heading 一并处理（标题不再漏）
    for i in range(1, 4):
        name = f"Heading {i}"
        if name in doc.styles:
            apply_font(doc.styles[name])

def strip_markdown_marks(s: str) -> str:
    """
    清理常见 Markdown 标记，避免 Word 里出现 ** ### ``` 等符号
    只做“去符号”，不做复杂排版转换（稳、够用）
    """
    if not s:
        return ""

    # 去掉代码块围栏
    s = re.sub(r"```.*?\n", "", s, flags=re.DOTALL)
    s = s.replace("```", "")

    # 标题符号
    s = re.sub(r"^\s{0,3}#{1,6}\s*", "", s, flags=re.MULTILINE)

    # 粗体/斜体标记
    s = s.replace("**", "")
    s = s.replace("__", "")
    s = s.replace("*", "")
    s = s.replace("_", "")

    # 行内代码
    s = s.replace("`", "")

    # 去掉多余空行
    s = re.sub(r"\n{3,}", "\n\n", s).strip()

    return s

def _pick_dim_scores_for_docx(band: dict):
    """
    兼容多种 band_scores 结构：
    1) {"Task Response": {"score": 6}, ...}
    2) {"Task Response": 6, ...}
    3) {"TR": 6, "CC": 6, "LR": 6, "GRA": 6, "overall": 6}
    4) 你的深度结构：{"TR":{"score":6},...,"overall":{"score":6.5,"avg":6.75}}
    """
    if not isinstance(band, dict):
        return None, None, None, None, None

    def as_score(v):
        if isinstance(v, dict) and isinstance(v.get("score"), (int, float)):
            return float(v["score"])
        if isinstance(v, (int, float)):
            return float(v)
        return None

    def pick(*names):
        for n in names:
            if n in band:
                sc = as_score(band.get(n))
                if sc is not None:
                    return sc
        return None

    TR = pick("Task Response", "Task Achievement", "TR")
    CC = pick("Coherence and Cohesion", "CC")
    LR = pick("Lexical Resource", "LR")
    GRA = pick("Grammatical Range & Accuracy", "Grammatical Range and Accuracy", "GRA")
    overall = pick("overall", "Overall", "OVERALL")

    if overall is None:
        parts = [x for x in (TR, CC, LR, GRA) if isinstance(x, (int, float))]
        overall = round((sum(parts) / len(parts)) * 2) / 2 if parts else None

    return TR, CC, LR, GRA, overall


def save_feedback_to_docx(
    feedback_text: str,
    output_path: str,
    task_type: Optional[str] = None,
    band_scores: Optional[Dict[str, Any]] = None,
    original_text: Optional[str] = None
) -> str:
    """
    生成 docx：
    1) 原文 vs 优化后对比表格
    2) 逐句批改反馈（原始文本）
    3) 段落优化（如果模型输出了；否则用修正句拼接）
    4) 优化依据（如果有）
    5) 评分建议（AI主裁评分，如果有）
    6) 【改错练习】（从【错误】提取）
    """
    def _clean_fix_line(s: str) -> str:
        s = (s or "").strip()
        # 去掉模型偶尔吐出的行首符号：全角/半角冒号、破折号、项目符号
        s = re.sub(r"^[：:\-\–\—•\u2022]+\s*", "", s)
        return s

    # ===== 预处理（函数体最开始）=====
    cleaned_feedback = strip_markdown_marks(feedback_text or "")
    if original_text:
        original_text = strip_markdown_marks(original_text)

    # 先用“完整反馈”解析出 optimized / rationale（不要先裁剪！）
    parsed = _parse_feedback_blocks(cleaned_feedback)
    items = parsed["items"]
    optimized = parsed["optimized"]
    rationale = parsed["rationale"]

    # 再做“原始输出区”专用版本：只保留逐句批改，去掉段落优化/优化依据（兼容带【】和不带【】）
    def _strip_opt_sections_for_raw(text: str) -> str:
        # 找到最早出现的“段落优化/优化依据”标题位置（任意一种写法），从那里截断
        patterns = [
            r"^\s*【\s*段落优化\s*】\s*[:：]?\s*$",
            r"^\s*段落优化\s*[:：]?\s*$",
            r"^\s*【\s*优化依据\s*】\s*[:：]?\s*$",
            r"^\s*优化依据\s*[:：]?\s*$",
        ]
        earliest = None
        for pat in patterns:
            m = re.search(pat, text, flags=re.MULTILINE)
            if m:
                pos = m.start()
                earliest = pos if earliest is None else min(earliest, pos)
        return text[:earliest].strip() if earliest is not None else text.strip()

    feedback_for_raw = _strip_opt_sections_for_raw(cleaned_feedback)

    # ===== 结束预处理 =====

    doc = Document()
    _set_doc_default_font(doc)

    # 注意：下面不要再重新 parsed/items/optimized/rationale 了
    # ---- 1) 标题（不再展示作文类型） ----
    doc.add_heading("作文批改反馈", level=1)

    # 反馈：直接放标题下方（用 rationale 的内容）
    doc.add_paragraph("反馈：")
    doc.add_paragraph(rationale if rationale else "无")
    
    # ---- 2) 原文 vs 优化后对比表格 ----
    if original_text:
        doc.add_heading("原文与优化后对比", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "原文"
        table.cell(0, 1).text = "优化后"

        table.cell(1, 0).text = original_text.strip()

        # ===== 关键改动开始：严格拼接【修正】，非英文/解析失败回退原文；补例子另起块 =====

        # ① 右侧“优化后”：逐句拼接，优先【修正】；遇到“非英文修正”则用 original
        final_lines = []
        if items and isinstance(items, list):
            for it in items:
                orig = (it.get("original") or "").strip()
                fix = (it.get("fix") or "").strip()

                # 【修正】为“无”或空 => 用原句
                if fix in ["无", "无。", "", "none", "None"]:
                    final_lines.append(orig)
                    continue

                # 先清洗，再判断
                clean_fix = _clean_fix_line(fix)

                # 清洗后为空 => 回退原句
                if not clean_fix:
                    final_lines.append(orig)
                    continue

                # 检测到“修正不是英文句子” => 回退原句
                if not _looks_english_sentence(clean_fix):
                    final_lines.append(orig)
                    continue

                # 正常情况：用清洗后的修正
                final_lines.append(clean_fix)

        final_for_table = "\n".join([x for x in final_lines if x.strip()]).strip()

        # ② 解析失败（items 为空 / 全部为空行）=> 回退原文
        if not final_for_table:
            final_for_table = (original_text or "").strip()

        table.cell(1, 1).text = final_for_table

        # ③ 第二次 DS 补例子：放在表格下面，让老师自己决定是否合并（不覆盖正文）
        optional_example = ""
        # 用于 Task 判定的输入：优先学生原文；若原文噪音太多/英文太少，则回退到 final_for_table
        task_detect_text = (original_text or "").strip()

        # 简单清洗：去掉明显中文行（可按你数据再调）
        task_detect_lines = []
        for ln in task_detect_text.splitlines():
            s = ln.strip()
            if not s:
                continue
            # 丢掉“纯中文/几乎全中文”的行
            cjk = len(re.findall(r"[\u4e00-\u9fff]", s))
            letters = len(re.findall(r"[A-Za-z]", s))
            if cjk >= 4 and letters < 5:
                continue
            task_detect_lines.append(s)

        task_detect_text = "\n".join(task_detect_lines).strip()

        # 如果清洗后英文太少，用 final_for_table（通常更干净）
        if len(re.findall(r"[A-Za-z]", task_detect_text)) < 30:
            task_detect_text = (final_for_table or "").strip()

        resolved_task_type = detect_task_type_via_deepseek(task_detect_text, api_key=None)

        if resolved_task_type == "Task 2":
            optional_example = generate_optional_example_with_deepseek(
                essay_text=final_for_table,
                task_type="Task 2",
                api_key=None,
                max_tokens=140,   # 更短一点，稳定输出 2-4 句
            )

        if optional_example:
            doc.add_paragraph("")
            doc.add_heading("可选补充示例", level=3)
            doc.add_paragraph("说明：以下仅为新增例子/细节本身；是否合并进正文由老师自行决定。")
            doc.add_paragraph(optional_example.strip())

        # ===== 关键改动结束 =====

    # ---- 3) 逐句反馈（原样保留，便于你对齐核对） ----
    doc.add_heading("逐句批改", level=2)
    doc.add_paragraph(feedback_for_raw.strip())

    # ✅ 注意：不再输出“段落优化”章节
    # ✅ 注意：不再输出“优化依据”章节（因为已经作为“反馈：”放到最上面了）

    # ---- 6) 评分建议（AI主裁评分） ----
    doc.add_heading("评分建议", level=2)

    TR, CC, LR, GRA, overall = _pick_dim_scores_for_docx(band_scores or {})

    if all(x is not None for x in (TR, CC, LR, GRA)):
        doc.add_paragraph(f"TR: {TR}")
        doc.add_paragraph(f"CC: {CC}")
        doc.add_paragraph(f"LR: {LR}")
        doc.add_paragraph(f"GRA: {GRA}")
        doc.add_paragraph(f"Overall（向下取 0.5）: {overall}")
    else:
        doc.add_paragraph("本次未获得 AI 评分（可能网络波动或服务端失败）。")

    # ---- 7) 改错练习（从【错误】提取） ----
    drills = _extract_error_drills(items)
    doc.add_heading("【改错练习】", level=2)
    if not drills:
        doc.add_paragraph("无")
    else:
        for d in drills:
            no = d.get("no")
            orig = d.get("original", "")
            err = d.get("error", "")
            doc.add_paragraph(f"{no}. 原句：{orig}")
            doc.add_paragraph(f"   错误：{err}")
            doc.add_paragraph("   正确句子：______________________________")

    doc.save(output_path)
    return output_path
