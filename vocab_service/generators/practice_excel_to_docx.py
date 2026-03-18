# generators/practice_excel_to_docx.py
# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Sequence, Tuple

from openpyxl import load_workbook
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_LINE_SPACING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell, Table

from core import field_definitions as fd


# =========================
# 固定样式（原样复制自 docx_generator.py）
# =========================
FONT_NAME = "等线"
FONT_SIZE = 11
TITLE_SIZE = 16

ROW_HEIGHT_HEADER_CM = 1.0
ROW_HEIGHT_BODY_CM = 2.0

# A4 边距（显式写死，避免 Word 默认边距导致“看起来列宽不对”）
MARGIN_LEFT_CM = 1.7
MARGIN_RIGHT_CM = 1.7
MARGIN_TOP_CM = 1.8
MARGIN_BOTTOM_CM = 1.8


# =========================
# docx 样式/固定列宽（关键：原样复制自 docx_generator.py）
# =========================
def _set_run_font(run, size_pt: int = FONT_SIZE, bold: Optional[bool] = None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold

def _set_cell_text(cell: _Cell, text: str, size_pt: int = FONT_SIZE, bold: Optional[bool] = None):
    cell.text = "" if text is None else str(text)

    for p in cell.paragraphs:
        # ✅ 关键：清掉段前段后，否则会把“1.0cm 行高”视觉上撑大
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = 1.0

        #（可选但安全）避免段落被额外拉开
        pf.keep_together = True

        if not p.runs:
            r = p.add_run("")
            _set_run_font(r, size_pt=size_pt, bold=bold)

        for r in p.runs:
            _set_run_font(r, size_pt=size_pt, bold=bold)

def _set_row_height(row, cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(int(Cm(cm).twips)))
    trHeight.set(qn("w:hRule"), "exact")

def _set_document_base_style(doc: Document):
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(FONT_SIZE)

    # Heading 2（用于标题）
    if "Heading 2" in styles:
        h2 = styles["Heading 2"]
        h2.font.name = FONT_NAME
        h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        h2.font.size = Pt(TITLE_SIZE)
        h2.font.bold = True

def _apply_page_setup(doc: Document, total_width_cm: float):
    """
    ✅ 练习版规则（按你最新口径锁死）：
    - 总列宽 > 19.5 -> 横版
    - 总列宽 <= 19.5 -> 竖版
    """
    section = doc.sections[0]
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)

    if total_width_cm > 19.5:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT

def _set_table_fixed_layout(table: Table):
    table.style = "Table Grid"
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr

    # tblLayout fixed
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

def _set_tbl_grid(table: Table, col_widths_cm: List[float]):
    """
    Word 最吃这一段：tblGrid
    """
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    else:
        for child in list(grid):
            grid.remove(child)

    for w_cm in col_widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(Cm(w_cm).twips)))
        grid.append(gridCol)

def _set_cell_width(cell: _Cell, w_cm: float):
    """
    写 tcW，防止 Word 对单元格再次“自适应”
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(Cm(w_cm).twips)))

def _apply_column_widths(table: Table, col_widths_cm: List[float]):
    _set_table_fixed_layout(table)
    _set_tbl_grid(table, col_widths_cm)

    # columns width + each cell tcW 双保险
    for i, w_cm in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w_cm)

    for row in table.rows:
        for i, w_cm in enumerate(col_widths_cm):
            _set_cell_width(row.cells[i], w_cm)

# =========================
# 修正单栏为双栏
# =========================
def _set_table_cell_margins(table: Table, top_cm=0.05, bottom_cm=0.05, left_cm=0.05, right_cm=0.05):
    """
    ✅ 进一步保证“行高紧凑”：缩小单元格内边距
    这是导致“看起来行高变大”的第二大元凶（第一是段前段后）。
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblPr.append(mar)

    def _set_side(tag, cm_val):
        el = mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            mar.append(el)
        el.set(qn("w:w"), str(int(Cm(cm_val).twips)))
        el.set(qn("w:type"), "dxa")

    _set_side("top", top_cm)
    _set_side("bottom", bottom_cm)
    _set_side("left", left_cm)
    _set_side("right", right_cm)


def _get_cell_value(row_values, idx: Dict[str, int], key: str) -> Any:
    col = idx.get(key)
    if col is None or col >= len(row_values):
        return None
    return row_values[col]


def _resolve_word_like_value(k: str, row_values, idx: Dict[str, int]) -> Any:
    if k == "word_original":
        return (
            _get_cell_value(row_values, idx, "word_original")
            or _get_cell_value(row_values, idx, "word")
            or _get_cell_value(row_values, idx, "英文单词")
            or _get_cell_value(row_values, idx, "word_norm")
            or _get_cell_value(row_values, idx, "word_display")
        )
    if k == "word_norm":
        return (
            _get_cell_value(row_values, idx, "word_norm")
            or _get_cell_value(row_values, idx, "word_original")
            or _get_cell_value(row_values, idx, "word_display")
        )
    return _get_cell_value(row_values, idx, k)


def _calc_text_for_cell(k: str, c: Dict[str, Any], row_values, idx: Dict[str, int], row_no: int) -> str:
    if k == "no":
        return str(row_no)
    if c.get("blank", False):
        return ""
    v = _resolve_word_like_value(k, row_values, idx)
    return "" if v is None else str(v)


def _build_table_double(
    doc: Document,
    ws,
    col_defs: List[Dict[str, Any]],
    col_widths_cm_double: List[float],
    body_row_height_cm: float,
):
    """
    双栏：每行放 2 个条目（左/右各一组 col_defs）
    字段映射完全复用 col_defs：不会写反“留空列”。
    """
    header_keys = _get_header_keys(ws)
    idx: Dict[str, int] = {k: i for i, k in enumerate(header_keys) if k}

    # 6 列：左3 + 右3
    table = doc.add_table(rows=1, cols=len(col_defs) * 2)
    _apply_column_widths(table, col_widths_cm_double)
    _set_table_cell_margins(table, top_cm=0.05, bottom_cm=0.05, left_cm=0.05, right_cm=0.05)

    # header row
    hdr = table.rows[0].cells
    for side in (0, 1):
        base = side * len(col_defs)
        for j, c in enumerate(col_defs):
            k = c["key"]
            w_cm = col_widths_cm_double[base + j]
            _set_cell_width(hdr[base + j], w_cm)
            _set_cell_text(hdr[base + j], _label_of(k), size_pt=FONT_SIZE, bold=True)

    _set_row_height(table.rows[0], ROW_HEIGHT_HEADER_CM)

    # body rows：两两配对
    rows_data = list(ws.iter_rows(min_row=2, values_only=True))
    n = len(rows_data)
    pairs = (n + 1) // 2

    for i in range(pairs):
        left_i = 2 * i
        right_i = 2 * i + 1

        row = table.add_row()
        cells = row.cells

        # 左侧：row_no = left_i + 1
        if left_i < n:
            r_left = rows_data[left_i]
            row_no_left = left_i + 1
            for j, c in enumerate(col_defs):
                k = c["key"]
                text = _calc_text_for_cell(k, c, r_left, idx, row_no_left)
                _set_cell_text(cells[j], text, size_pt=FONT_SIZE, bold=False)
                _set_cell_width(cells[j], col_widths_cm_double[j])
        else:
            for j in range(len(col_defs)):
                _set_cell_text(cells[j], "", size_pt=FONT_SIZE, bold=False)
                _set_cell_width(cells[j], col_widths_cm_double[j])

        # 右侧：row_no = right_i + 1
        base = len(col_defs)
        if right_i < n:
            r_right = rows_data[right_i]
            row_no_right = right_i + 1
            for j, c in enumerate(col_defs):
                k = c["key"]
                text = _calc_text_for_cell(k, c, r_right, idx, row_no_right)
                _set_cell_text(cells[base + j], text, size_pt=FONT_SIZE, bold=False)
                _set_cell_width(cells[base + j], col_widths_cm_double[base + j])
        else:
            for j in range(len(col_defs)):
                _set_cell_text(cells[base + j], "", size_pt=FONT_SIZE, bold=False)
                _set_cell_width(cells[base + j], col_widths_cm_double[base + j])

        _set_row_height(row, cm=body_row_height_cm)

    return table

# =========================
# practice：只做“视图 + 留空列”
# =========================
def _stamp(ts: str, run_id: str) -> str:
    ts2 = ts.strip()
    rid = run_id.strip()
    if rid:
        return f"{ts2}_{rid}"
    return ts2

def _get_header_keys(ws) -> List[str]:
    header = [c.value for c in ws[1]]
    out: List[str] = []
    for h in header:
        if h is None:
            out.append("")
        else:
            out.append(str(h).strip())
    return out

def _label_of(key: str) -> str:
    # 表头显示 label（来自 field_definitions）
    if key in fd.FIELD_DEFS:
        return fd.FIELD_DEFS[key].label
    return key

def _sheet_has_required_keys(ws, col_defs: List[Dict[str, Any]]) -> Tuple[bool, List[str]]:
    header_keys = _get_header_keys(ws)
    missing: List[str] = []
    for c in col_defs:
        if c.get("blank", False):
            continue
        k = c["key"]
        if k not in header_keys:
            missing.append(k)
    return (len(missing) == 0, missing)

def _build_one_docx_from_workbook(
    wb,
    output_path: str,
    doc_title: str,
    col_defs: List[Dict[str, Any]],
    col_widths_cm: List[float],
    body_row_height_cm: float,
    layout: str = "single",  # ✅ 新增：single / double
):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()
    _set_document_base_style(doc)

    # 标题（可删：如果你希望练习 docx 没有标题，就把这一段注释掉）
    p = doc.add_heading(doc_title, level=2)
    for r in p.runs:
        _set_run_font(r, size_pt=TITLE_SIZE, bold=True)

    total_w = sum(col_widths_cm)
    _apply_page_setup(doc, total_w)

    # 每个 sheet -> 一张表
    for sheet in wb.sheetnames:
        ws = wb[sheet]

        ok, missing = _sheet_has_required_keys(ws, col_defs)
        if not ok:
            # 如果某 sheet 缺字段：直接报错（更符合“契约一致”）
            raise ValueError(f"[practice_docx] sheet={sheet} missing columns: {missing}")

        header_keys = _get_header_keys(ws)
        idx: Dict[str, int] = {k: i for i, k in enumerate(header_keys) if k}

        if layout == "double":
            # ✅ 双栏：col_widths_cm 必须是 6 列宽
            _build_table_double(
                doc=doc,
                ws=ws,
                col_defs=col_defs,
                col_widths_cm_double=col_widths_cm,
                body_row_height_cm=body_row_height_cm,
            )
        else:
            # ✅ 单栏：保留你原逻辑（仅加一个 cell margins，让行高更贴近 1.0cm）
            table = doc.add_table(rows=1, cols=len(col_defs))
            _apply_column_widths(table, col_widths_cm)
            _set_table_cell_margins(table, top_cm=0.05, bottom_cm=0.05, left_cm=0.05, right_cm=0.05)

            # header row
            hdr = table.rows[0].cells
            for j, c in enumerate(col_defs):
                k = c["key"]
                w_cm = col_widths_cm[j]
                _set_cell_width(hdr[j], w_cm)
                _set_cell_text(hdr[j], _label_of(k), size_pt=FONT_SIZE, bold=True)
            _set_row_height(table.rows[0], ROW_HEIGHT_HEADER_CM)

            # body rows
            for row_no, r in enumerate(ws.iter_rows(min_row=2, values_only=True), start=1):
                row = table.add_row()
                cells = row.cells
                for j, c in enumerate(col_defs):
                    k = c["key"]
                    text = _calc_text_for_cell(k, c, r, idx, row_no)

                    _set_cell_text(cells[j], text, size_pt=FONT_SIZE, bold=False)
                    _set_cell_width(cells[j], col_widths_cm[j])

                _set_row_height(row, cm=body_row_height_cm)

        doc.add_page_break()

    doc.save(output_path)
    print(f"[DONE] {output_path}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input_xlsx", required=True, help="shuffle 后的 master xlsx（表头为 key）")
    parser.add_argument("--output_dir", required=True, help="输出目录")
    parser.add_argument("--run_id", default="", help="可选：run_id（例如 run01）")
    parser.add_argument("--timestamp", default="", help="可选：timestamp（默认当前时间）")
    args = parser.parse_args()

    ts = args.timestamp.strip() or datetime.now().strftime("%Y%m%d_%H%M%S")
    suffix = _stamp(ts, args.run_id)

    in_xlsx = args.input_xlsx
    out_dir = args.output_dir
    os.makedirs(out_dir, exist_ok=True)

    wb = load_workbook(in_xlsx, data_only=True)

    # =========================
    # 4 类练习（学生版）
    # 只允许你调整：col_widths_cm 和 body_row_height_cm
    # =========================
    DOUBLE_WORD_COL_WIDTHS_CM = [1.0, 3.8, 3.8, 1.0, 3.8, 3.8]

    # 单词 英译中：给英文单词；中文解释留空让学生写
    spec_word_e2c = {
        "title": "单词 英译中（学生版）",
        "filename": f"练习_单词_英译中_{suffix}.docx",
        "col_defs": [
            {"key": "no", "blank": False},
            {"key": "word_original", "blank": False},
            {"key": "pos_cn", "blank": True},   # ✅ 留空
        ],
        "layout": "double",
        "col_widths_cm": DOUBLE_WORD_COL_WIDTHS_CM,
        "body_row_height_cm": 1.0,             # 按你要求
    }

    # 单词 中译英：给中文解释；英文单词留空让学生写
    spec_word_c2e = {
        "title": "单词 中译英（学生版）",
        "filename": f"练习_单词_中译英_{suffix}.docx",
        "col_defs": [
            {"key": "no", "blank": False},
            {"key": "pos_cn", "blank": False},
            {"key": "word_original", "blank": True},  # ✅ 留空
        ],
        "layout": "double",
        "col_widths_cm": DOUBLE_WORD_COL_WIDTHS_CM,
        "body_row_height_cm": 1.0,
    }

    # 例句 英译中：给英文例句；例句翻译留空让学生写
    spec_sent_e2c = {
        "title": "例句 英译中（学生版）",
        "filename": f"练习_例句_英译中_{suffix}.docx",
        "col_defs": [
            {"key": "no", "blank": False},
            {"key": "example", "blank": False},
            {"key": "example_cn", "blank": True},  # ✅ 留空
        ],
        "layout": "single",
        "col_widths_cm": [1.2, 9.0, 9.0],
        "body_row_height_cm": 2.0,  # 按你要求
    }

    # 例句 中译英：给例句翻译；英文例句留空让学生写
    spec_sent_c2e = {
        "title": "例句 中译英（学生版）",
        "filename": f"练习_例句_中译英_{suffix}.docx",
        "col_defs": [
            {"key": "no", "blank": False},
            {"key": "example_cn", "blank": False},
            {"key": "example", "blank": True},  # ✅ 留空
        ],
        "layout": "single",
        "col_widths_cm": [1.2, 9.0, 9.0],
        "body_row_height_cm": 2.0,
    }

    specs = [spec_word_e2c, spec_word_c2e, spec_sent_e2c, spec_sent_c2e]

    # 如果某类练习需要的字段在 Excel 根本不存在：跳过并提示（符合你“前端不展示缺失选项”的方向）
    # 判断方式：只看第一个 sheet 的 header 是否具备 required key
    first_ws = wb[wb.sheetnames[0]]
    header_keys = set(_get_header_keys(first_ws))

    for sp in specs:
        required = [c["key"] for c in sp["col_defs"] if not c.get("blank", False)]
        missing_req = [k for k in required if k not in header_keys]
        if missing_req:
            print(f"[SKIP] {sp['filename']} missing required columns in excel: {missing_req}")
            continue

        out_path = os.path.join(out_dir, sp["filename"])
        _build_one_docx_from_workbook(
            wb=wb,
            output_path=out_path,
            doc_title=sp["title"],
            col_defs=sp["col_defs"],
            col_widths_cm=sp["col_widths_cm"],
            body_row_height_cm=sp["body_row_height_cm"],
            layout=sp.get("layout", "single"),
        )

    print(f"[DONE] wrote practice docx to: {os.path.abspath(out_dir)}")


if __name__ == "__main__":
    main()
