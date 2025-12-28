# generators/docx_generator.py
# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import os
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, List, Optional, Sequence, Tuple

from openpyxl import load_workbook
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell, Table


# =========================
# 固定样式（按你的需求）
# =========================
FONT_NAME = "等线"
FONT_SIZE = 11
TITLE_SIZE = 16

ROW_HEIGHT_HEADER_CM = 1.0
ROW_HEIGHT_BODY_CM = 2.0

# A4 边距（显式写死，避免 Word 默认边距导致“看起来列宽不对”）
MARGIN_LEFT_CM = 1.7
MARGIN_RIGHT_CM = 1.7
MARGIN_TOP_CM = 1.8
MARGIN_BOTTOM_CM = 1.8


# =========================
# 数据结构
# =========================
@dataclass
class ExcelRow:
    no: int
    word_original: str
    phonetic_uk: str
    pos_cn: str          # 中文解释（契约：表头=中文解释）
    definition_en: str   # 英文解释（契约：表头=英文解释）
    example: str
    synonyms: str


# =========================
# Excel 读取
# =========================
def _norm_header(s: str) -> str:
    # 统一：去空格/连字符/全角符号差异
    s = (s or "").strip().lower()
    s = s.replace("＋", "+")
    return re.sub(r"[\s\-]+", "_", s)

def _find_col(header: Sequence[Any], targets: Sequence[str]) -> Optional[int]:
    norm = [_norm_header(str(x)) for x in header]
    for t in targets:
        t2 = _norm_header(t)
        if t2 in norm:
            return norm.index(t2)
    return None

def _safe_int(v: Any, default: int) -> int:
    try:
        if v is None:
            return default
        if isinstance(v, (int, float)):
            return int(v)
        s = str(v).strip()
        if not s:
            return default
        return int(float(s))
    except Exception:
        return default

def load_rows_from_vocab_excel(xlsx_path: str, only_sheet: Optional[str] = None) -> List[ExcelRow]:
    """
    读取“词汇笔记.xlsx”（完整词汇表视图）。
    必须遵守列名契约（输出侧）：
      - 中文解释
      - 英文解释

    读入侧允许历史兼容（仅为避免旧文件导致空列）：
      - 词性+中文 / 词性＋中文
    """
    wb = load_workbook(xlsx_path, data_only=True)
    sheet_name = only_sheet or wb.sheetnames[0]
    ws = wb[sheet_name]
    all_rows = list(ws.iter_rows(values_only=True))
    if not all_rows:
        return []

    header = list(all_rows[0])

    idx_no = _find_col(header, ["no", "序号", "编号"])
    idx_word = _find_col(header, ["word_original", "word", "英文单词", "英文词汇"])
    idx_uk = _find_col(header, ["phonetic_uk", "英式音标", "英式发音", "英音音标"])

    # 核心契约：中文解释 / 英文解释
    idx_pos = _find_col(
        header,
        ["pos_cn", "中文解释", "词性+中文", "词性＋中文", "词性中文"]  # 历史兼容只读
    )
    idx_def_en = _find_col(
        header,
        ["definition_en", "英文解释", "英文释义", "英文定义", "英文意思"]
    )

    idx_ex = _find_col(header, ["example", "例句"])
    idx_syn = _find_col(header, ["synonyms", "同义替换", "同义词", "近义替换"])

    if idx_word is None:
        raise ValueError(f"Excel missing required column: 英文单词/word_original. file={xlsx_path}")

    rows: List[ExcelRow] = []
    for r in all_rows[1:]:
        word = str(r[idx_word] or "").strip()
        if not word:
            continue

        no = _safe_int(r[idx_no], default=len(rows) + 1) if idx_no is not None else (len(rows) + 1)
        uk = str(r[idx_uk] or "").strip() if idx_uk is not None else ""
        pos_cn = str(r[idx_pos] or "").strip() if idx_pos is not None else ""
        def_en = str(r[idx_def_en] or "").strip() if idx_def_en is not None else ""
        ex = str(r[idx_ex] or "").strip() if idx_ex is not None else ""
        syn = str(r[idx_syn] or "").strip() if idx_syn is not None else ""

        rows.append(ExcelRow(
            no=no,
            word_original=word,
            phonetic_uk=uk,
            pos_cn=pos_cn,
            definition_en=def_en,
            example=ex,
            synonyms=syn,
        ))
    return rows


# =========================
# docx 样式/固定列宽（关键）
# =========================
def _set_run_font(run, size_pt: int = FONT_SIZE, bold: Optional[bool] = None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold

def _set_cell_text(cell: _Cell, text: str, size_pt: int = FONT_SIZE, bold: Optional[bool] = None):
    cell.text = "" if text is None else str(text)
    for p in cell.paragraphs:
        if not p.runs:
            r = p.add_run("")
            _set_run_font(r, size_pt=size_pt, bold=bold)
        for r in p.runs:
            _set_run_font(r, size_pt=size_pt, bold=bold)

def _set_row_height(row, cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(int(Cm(cm).twips)))
    trHeight.set(qn("w:hRule"), "exact")

def _set_document_base_style(doc: Document):
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(FONT_SIZE)

    # Heading 2（用于标题）
    if "Heading 2" in styles:
        h2 = styles["Heading 2"]
        h2.font.name = FONT_NAME
        h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        h2.font.size = Pt(TITLE_SIZE)
        h2.font.bold = True

def _apply_page_setup(doc: Document, total_width_cm: float):
    """
    需求案规则：
    - 总列宽 >= 19.5 -> 横版
    - < 19.5 -> 竖版
    """
    section = doc.sections[0]
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)

    if total_width_cm >= 19.5:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT

def _set_table_fixed_layout(table: Table):
    table.style = "Table Grid"
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr

    # tblLayout fixed
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

def _set_tbl_grid(table: Table, col_widths_cm: List[float]):
    """
    Word 最吃这一段：tblGrid
    """
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    else:
        for child in list(grid):
            grid.remove(child)

    for w_cm in col_widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(Cm(w_cm).twips)))
        grid.append(gridCol)

def _set_cell_width(cell: _Cell, w_cm: float):
    """
    写 tcW，防止 Word 对单元格再次“自适应”
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(Cm(w_cm).twips)))

def _apply_column_widths(table: Table, col_widths_cm: List[float]):
    _set_table_fixed_layout(table)
    _set_tbl_grid(table, col_widths_cm)

    # columns width + each cell tcW 双保险
    for i, w_cm in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w_cm)

    for row in table.rows:
        for i, w_cm in enumerate(col_widths_cm):
            _set_cell_width(row.cells[i], w_cm)


# =========================
# docx 生成：词汇笔记
# =========================
def build_vocab_note_docx(
    rows: Optional[List[ExcelRow]] = None,
    output_path: str = "",
    variant: str = "all_fields",
    *,
    entries: Optional[List[Any]] = None,
    xlsx_path: Optional[str] = None,
    only_sheet: Optional[str] = None,
):
    """
    兼容版 build_vocab_note_docx

    支持三种输入方式（按优先级）：
    1) entries=...   （API export 直接传 run_cache["entries"]，元素可以是 dict 或 Entry 对象）
    2) rows=...      （旧用法：List[ExcelRow]）
    3) xlsx_path=... （从你导出的 vocab_note_excel 读 Excel 再生成 docx）

    旧代码兼容：
      build_vocab_note_docx(vocab_rows, "xxx.docx", variant="all_fields")
    新代码用法：
      build_vocab_note_docx(output_path="xxx.docx", variant="all_fields", entries=entries)
    """

    # --------- 1) 统一拿到 rows(List[ExcelRow]) ----------
    def _to_text(v: Any) -> str:
        if v is None:
            return ""
        if isinstance(v, (list, tuple)):
            # synonyms 可能是 list
            return "; ".join([str(x).strip() for x in v if str(x).strip()])
        return str(v).strip()

    def _entry_to_excel_row(e: Any, no_default: int) -> ExcelRow:
        # e 可能是 dict / dataclass Entry / 其它对象
        if isinstance(e, dict):
            d = e
        else:
            d = getattr(e, "__dict__", {}) or {}

        no = _safe_int(d.get("no"), default=no_default)
        word_original = _to_text(
            d.get("word_display")
            or d.get("word_norm")
            or d.get("word_original")
            or d.get("word")
            or d.get("英文单词")
        )

        phonetic_uk = _to_text(d.get("phonetic_uk") or d.get("英式音标"))

        # 这两个字段在你的系统契约里就是 pos_cn / definition_en
        # 如果历史数据里还是“词性+中文”，这里也做兼容兜底
        pos_cn = _to_text(d.get("pos_cn") or d.get("中文解释") or d.get("词性+中文") or d.get("词性＋中文"))
        definition_en = _to_text(d.get("definition_en") or d.get("英文解释") or d.get("英文释义") or d.get("英文定义"))

        example = _to_text(d.get("example") or d.get("例句"))
        synonyms = d.get("synonyms")
        # synonyms 在 Entry 里是 List[str]，在 ExcelRow 里是 str
        synonyms_text = _to_text(synonyms)

        return ExcelRow(
            no=no,
            word_original=word_original,
            phonetic_uk=phonetic_uk,
            pos_cn=pos_cn,
            definition_en=definition_en,
            example=example,
            synonyms=synonyms_text,
        )

    if entries is not None:
        rows2: List[ExcelRow] = []
        for i, e in enumerate(entries, start=1):
            r = _entry_to_excel_row(e, no_default=i)
            # 过滤掉空单词行
            if r.word_original.strip():
                rows2.append(r)
        rows = rows2

    if rows is None and xlsx_path:
        rows = load_rows_from_vocab_excel(xlsx_path, only_sheet=only_sheet)

    if not rows:
        raise ValueError("build_vocab_note_docx: no rows/entries to write")

    if not output_path:
        raise ValueError("build_vocab_note_docx: output_path is required")

    # --------- 2) 以下保持你原来的 docx 生成逻辑不变 ----------
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    _set_document_base_style(doc)

    p = doc.add_heading("词汇笔记", level=2)
    for r in p.runs:
        _set_run_font(r, size_pt=TITLE_SIZE, bold=True)

    cols: List[Tuple[str, str, float]] = []

    if variant == "all_fields":
        cols = [
            ("word_original", "英文单词", 3.0),
            ("phonetic_uk", "英式音标", 3.0),
            ("pos_cn", "中文解释", 4.0),
            ("example", "例句", 9.5),
            ("synonyms", "同义替换", 6.0),
        ]
    elif variant == "with_example":
        cols = [
            ("no", "No.", 1.2),
            ("word_original", "英文单词", 4.0),
            ("phonetic_uk", "英式音标", 4.0),
            ("pos_cn", "中文解释", 5.0),
            ("example", "例句", 10.0),
        ]
    elif variant == "with_synonyms":
        cols = [
            ("no", "No.", 1.2),
            ("word_original", "英文单词", 3.0),
            ("phonetic_uk", "英式音标", 3.3),
            ("pos_cn", "中文解释", 4.0),
            ("synonyms", "同义替换", 7.0),
        ]
    elif variant == "basic_cn":
        cols = [
            ("no", "No.", 1.2),
            ("word_original", "英文单词", 4.0),
            ("phonetic_uk", "英式音标", 4.0),
            ("pos_cn", "中文解释", 8.0),
        ]
    else:
        raise ValueError(f"Unknown variant: {variant}")

    col_widths = [w for _, _, w in cols]
    total_w = sum(col_widths)
    _apply_page_setup(doc, total_w)

    table = doc.add_table(rows=1, cols=len(cols))
    _apply_column_widths(table, col_widths)

    hdr = table.rows[0].cells
    for j, (_, title, w_cm) in enumerate(cols):
        _set_cell_width(hdr[j], w_cm)
        _set_cell_text(hdr[j], title, size_pt=FONT_SIZE, bold=True)
    _set_row_height(table.rows[0], ROW_HEIGHT_HEADER_CM)

    for r in rows:
        tr = table.add_row()
        _set_row_height(tr, ROW_HEIGHT_BODY_CM)
        for j, (key, _, w_cm) in enumerate(cols):
            c = tr.cells[j]
            _set_cell_width(c, w_cm)
            val = getattr(r, key, "")
            _set_cell_text(c, val, size_pt=FONT_SIZE, bold=False)

    doc.save(output_path)
    
# =========================
# CLI
# =========================
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--vocab_excel", required=True)
    parser.add_argument("--out_dir", required=True)
    parser.add_argument("--vocab_sheet", default="")
    parser.add_argument("--run_id", default="")
    parser.add_argument("--timestamp", default="")
    args = parser.parse_args()

    ts = args.timestamp.strip() or datetime.now().strftime("%Y%m%d_%H%M%S")
    rid = args.run_id.strip()
    suffix = f"{ts}_{rid}" if rid else ts

    out_dir = args.out_dir
    os.makedirs(out_dir, exist_ok=True)

    vocab_rows = load_rows_from_vocab_excel(args.vocab_excel, only_sheet=(args.vocab_sheet.strip() or None))
    if not vocab_rows:
        raise SystemExit("[ERROR] vocab_excel has no rows")

    build_vocab_note_docx(vocab_rows, os.path.join(out_dir, f"词汇笔记_{suffix}_全字段.docx"), variant="all_fields")
    build_vocab_note_docx(vocab_rows, os.path.join(out_dir, f"词汇笔记_{suffix}_包含例句.docx"), variant="with_example")
    build_vocab_note_docx(vocab_rows, os.path.join(out_dir, f"词汇笔记_{suffix}_包含同义替换.docx"), variant="with_synonyms")
    build_vocab_note_docx(vocab_rows, os.path.join(out_dir, f"词汇笔记_{suffix}_基础中英释义.docx"), variant="basic_cn")

    print(f"[DONE] wrote docx to: {os.path.abspath(out_dir)}")


if __name__ == "__main__":
    main()