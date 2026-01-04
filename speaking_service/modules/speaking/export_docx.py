# export_docx.py
# Version B: QA 教学版（题目 + Outline + Model Answer）
# 输入：runs/<run_id>/answers/ 下的 Sxx_P1_*.txt 与 Sxx_P23_*.txt（来自 run_generate_answers.py）
# 输出：runs/<run_id>/export_docx/ 下的 docx
#
# 支持：
# - 每 segment 一个 docx（默认）
# - merge_all=true 合并总讲义（按 segment_id 顺序）
#
# 重要：不改你现有的字段、函数名、缓存逻辑；只做“导出层”。

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Tuple

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from dataclasses import dataclass
from typing import Literal

# -------------------------
# 固定 Header（按你要求）
# -------------------------

EXAM_NAME = "IELTS Speaking"
DEFAULT_BAND = "6.5-7"
DEFAULT_STYLE = "Natural & Native"
DEFAULT_ROLE = "High School Student (Hangzhou, China)"

SHOW_HEADER = False   # 是否在 docx 中展示 Header 区

PageBreakMode = Literal["none", "part", "segment"]

@dataclass
class ExportOptions:
    """
    统一的导出开关（给 Version A/B/C 共用）
    - show_header: 是否输出 Header(KV 区块)。注意：不等于 Topic 标题（标题建议一直保留）。
    - show_outline: 是否输出 Outline（仅 B/C 可能用到）
    - show_answers: 是否输出 Answer（仅 B/C 可能用到）
    - page_break:
        - "none"    : 不强制分页
        - "part"    : Part 之间分页（同一 segment 内）
        - "segment" : segment 之间分页（merge_all 的合并文档最常用）
    """
    show_header: bool = False
    show_outline: bool = True
    show_answers: bool = True
    page_break: PageBreakMode = "none"

# Header 开关：分别控制 A / B / C

SHOW_HEADER_A = False   # Version A（纯题目版）一般不需要
SHOW_HEADER_B = False   # Version B（教学版）可以选择不需要
SHOW_HEADER_C = False   # Version C（宣传截图版）一般不需要


# -------------------------
# 数据结构
# -------------------------

@dataclass
class QAItem:
    question: str
    outline: str
    answer: str


@dataclass
class Part2Block:
    cue: str
    bullets: List[str] = field(default_factory=list)
    outline: str = ""
    answer: str = ""


@dataclass
class SegmentDoc:
    segment_id: int
    topic: str
    part1: List[QAItem] = field(default_factory=list)
    part2: Optional[Part2Block] = None
    part3: List[QAItem] = field(default_factory=list)


# -------------------------
# 工具：文件名安全
# -------------------------

def maybe_page_break(doc: Document, opts: "ExportOptions"):
    # opts.page_break: "none" | "segment" | "part"（你自己定义的枚举也行）
    if not opts:
        doc.add_page_break()
        return
    mode = getattr(opts, "page_break", None) or getattr(opts, "page_break_mode", None)
    if str(mode).lower() == "none":
        return
    doc.add_page_break()


def safe_filename(s: str) -> str:
    s = (s or "").strip()
    for ch in r'\/:*?"<>|':
        s = s.replace(ch, "_")
    s = re.sub(r"\s+", " ", s).strip()
    return s[:60] if len(s) > 60 else (s or "UnknownTopic")

# ✅ 向后兼容旧代码 / 不同版本调用
_safe_filename = safe_filename

def ensure_dir(p: str) -> str:
    os.makedirs(p, exist_ok=True)
    return p

import re

def _safe_filename(name: str, max_len: int = 80) -> str:
    """
    将 topic 变成 Windows 安全文件名：去除非法字符、压缩空白、限制长度。
    """
    if name is None:
        return "Untitled"
    s = str(name).strip()
    if not s:
        return "Untitled"

    # Windows 禁止字符: \ / : * ? " < > |
    s = re.sub(r'[\\/:*?"<>|]+', "_", s)
    # 控制字符
    s = re.sub(r"[\x00-\x1f]+", "", s)
    # 多空白->单下划线
    s = re.sub(r"\s+", "_", s).strip("_")

    if len(s) > max_len:
        s = s[:max_len].rstrip("_")
    return s or "Untitled"

# -------------------------
# 解析 answers txt（来自 render_part1 / render_part23）
# -------------------------

def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def parse_part1_file(path: str) -> SegmentDoc:
    """
    解析 Sxx_P1_*.txt
    预期结构（你当前 run_generate_answers.py 输出）大致：
      SEGMENT: 1
      TOPIC: xxx
      PART: 1

      Q1: ...
      Outline:
      ...
      Full Answer:
      ...
    """
    text = _read_text(path)
    seg_m = re.search(r"SEGMENT:\s*(\d+)", text)
    topic_m = re.search(r"TOPIC:\s*(.+)", text)

    seg_id = int(seg_m.group(1)) if seg_m else 0
    topic = topic_m.group(1).strip() if topic_m else "UnknownTopic"

    seg = SegmentDoc(segment_id=seg_id, topic=topic)

    # 分块：每个 Qx:
    blocks = re.split(r"\nQ\d+:\s*", text)
    # blocks[0] 是头部，后面每个 block 开头是 question 文本
    for b in blocks[1:]:
        # question 在第一行到换行
        lines = b.splitlines()
        if not lines:
            continue
        question = lines[0].strip()

        # Outline / Full Answer 抓取
        outline = ""
        answer = ""

        # 用标记定位
        out_m = re.search(r"Outline:\s*\n(.*?)\nFull Answer:\s*\n(.*)", b, flags=re.S)
        if out_m:
            outline = out_m.group(1).strip()
            answer = out_m.group(2).strip()
        else:
            # 兜底：尽量找 Outline: 之后到 Full Answer:
            o2 = re.search(r"Outline:\s*\n(.*?)(?:\nFull Answer:\s*\n|$)", b, flags=re.S)
            a2 = re.search(r"Full Answer:\s*\n(.*)$", b, flags=re.S)
            if o2:
                outline = o2.group(1).strip()
            if a2:
                answer = a2.group(1).strip()

        seg.part1.append(QAItem(question=question, outline=outline, answer=answer))

    return seg


def parse_part23_file(path: str) -> SegmentDoc:
    """
    解析 Sxx_P23_*.txt
    预期结构：
      SEGMENT: 1
      TOPIC: xxx

      PART: 2
      CUE: ...
      BULLET: ...
      Outline:
      ...
      Full Answer:
      ...

      PART: 3
      Q1: ...
      Outline:
      ...
      Full Answer:
      ...
    """
    text = _read_text(path)
    seg_m = re.search(r"SEGMENT:\s*(\d+)", text)
    topic_m = re.search(r"TOPIC:\s*(.+)", text)

    seg_id = int(seg_m.group(1)) if seg_m else 0
    topic = topic_m.group(1).strip() if topic_m else "UnknownTopic"
    seg = SegmentDoc(segment_id=seg_id, topic=topic)

    # Part 2
    cue = ""
    bullets: List[str] = []

    p2_block_m = re.search(r"PART:\s*2\s*\n(.*?)(?:\nPART:\s*3|\Z)", text, flags=re.S)
    p2_block = p2_block_m.group(1) if p2_block_m else ""

    # 在 PART 2 块里按行找 CUE
    cue = ""
    cue_m = re.search(r"^CUE:\s*(.+)$", p2_block, flags=re.M)
    if cue_m:
        cue = cue_m.group(1).strip()


    for bm in re.finditer(r"^BULLET:\s*(.+)$", text, flags=re.M):
        bullets.append(bm.group(1).strip())

    # Part2 outline/answer
    p2_m = re.search(r"PART:\s*2.*?Outline:\s*\n(.*?)\nFull Answer:\s*\n(.*?)(?:\nPART:\s*3|\Z)", text, flags=re.S)
    if p2_m:
        outline2 = p2_m.group(1).strip()
        answer2 = p2_m.group(2).strip()
        seg.part2 = Part2Block(cue=cue, bullets=bullets, outline=outline2, answer=answer2)
    else:
        # 如果没有 Part2 也允许 Part3-only
        if cue or bullets:
            seg.part2 = Part2Block(cue=cue, bullets=bullets)

    # Part 3 questions blocks
    p3_text_m = re.search(r"PART:\s*3\s*\n(.*)$", text, flags=re.S)
    if p3_text_m:
        p3_text = p3_text_m.group(1)
        blocks = re.split(r"\nQ\d+:\s*", "\n" + p3_text)
        for b in blocks[1:]:
            lines = b.splitlines()
            if not lines:
                continue
            question = lines[0].strip()

            outline = ""
            answer = ""
            out_m = re.search(r"Outline:\s*\n(.*?)\nFull Answer:\s*\n(.*)", b, flags=re.S)
            if out_m:
                outline = out_m.group(1).strip()
                answer = out_m.group(2).strip()
            else:
                o2 = re.search(r"Outline:\s*\n(.*?)(?:\nFull Answer:\s*\n|$)", b, flags=re.S)
                a2 = re.search(r"Full Answer:\s*\n(.*)$", b, flags=re.S)
                if o2:
                    outline = o2.group(1).strip()
                if a2:
                    answer = a2.group(1).strip()

            seg.part3.append(QAItem(question=question, outline=outline, answer=answer))

    return seg


def load_segments_from_answers_dir(answers_dir: str) -> Dict[int, SegmentDoc]:
    """
    从 answers_dir 读入所有 segment。
    允许 segment 同时有 P1 与 P23 文件，合并为一个 SegmentDoc。
    """
    seg_map: Dict[int, SegmentDoc] = {}

    if not os.path.isdir(answers_dir):
        raise FileNotFoundError(f"answers_dir not found: {answers_dir}")

    files = [f for f in os.listdir(answers_dir) if f.lower().endswith(".txt")]

    # 先按文件名粗分类
    p1_files = [f for f in files if "_P1_" in f]
    p23_files = [f for f in files if "_P23_" in f]

    for fn in p1_files:
        path = os.path.join(answers_dir, fn)
        seg = parse_part1_file(path)
        sid = seg.segment_id
        if sid not in seg_map:
            seg_map[sid] = seg
        else:
            # 合并
            seg_map[sid].topic = seg_map[sid].topic or seg.topic
            seg_map[sid].part1 = seg.part1

    for fn in p23_files:
        path = os.path.join(answers_dir, fn)
        seg = parse_part23_file(path)
        sid = seg.segment_id
        if sid not in seg_map:
            seg_map[sid] = seg
        else:
            seg_map[sid].topic = seg_map[sid].topic or seg.topic
            seg_map[sid].part2 = seg.part2
            seg_map[sid].part3 = seg.part3

    return seg_map


# -------------------------
# docx 渲染（Version B）
# -------------------------

def _set_doc_style(doc: Document):
    """
    统一教学讲义的基础样式（只管“好看”，不改业务逻辑）
    """
    def _set_style_font(style_obj, font_name: str, size_pt: int, bold: bool = False):
        # python-docx 的 style.font.name 只覆盖一部分；
        # 为了让英文/数字/中文都稳定变成“等线”，这里把 rFonts 全部写齐
        style_obj.font.name = font_name
        rFonts = style_obj._element.rPr.rFonts
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)
        style_obj.font.size = Pt(size_pt)
        style_obj.font.bold = bold

    # Normal
    normal = doc.styles["Normal"]
    _set_style_font(normal, "等线", 11, bold=False)

    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # ✅ Heading 1（V-C 的 PART 1/2/3 — topic 就是 level=1）
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        _set_style_font(h1, "等线", 16, bold=True)
        h1_pf = h1.paragraph_format
        h1_pf.space_before = Pt(12)
        h1_pf.space_after = Pt(6)

    # （可选但建议）Heading 2：你 V-A/V-B 里也用到
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        _set_style_font(h2, "等线", 14, bold=True)
        h2_pf = h2.paragraph_format
        h2_pf.space_before = Pt(10)
        h2_pf.space_after = Pt(4)

# 分隔线

def _add_divider(doc: Document):
    """分隔线（比空行更讲义）"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)

    # 用段落底边框实现横线
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

# 分隔线
def _add_kv(doc: Document, k: str, v: str):
    """Header 的 Key: Value（Key 加粗）"""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{k}: ")
    r1.bold = True
    p.add_run(v or "")

# 分隔线
def _add_bullets(doc: Document, items: List[str]):
    """把多行变成 bullet list"""
    items = [x.strip("•- \t").strip() for x in (items or []) if str(x).strip()]
    if not items:
        p = doc.add_paragraph("(empty)")
        p.runs[0].italic = True
        return
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

# 分隔线
def _add_answer_paragraphs(doc: Document, text: str):
    """
    Model Answer：尽量变成“正常段落”，避免一行一行像日志
    - 连续空行视为分段
    """
    t = (text or "").strip()
    if not t:
        p = doc.add_paragraph("(empty)")
        p.runs[0].italic = True
        return

    # 以空行分段
    paras = [p.strip() for p in re.split(r"\n\s*\n", t) if p.strip()]
    for para in paras:
        # 把段内的换行压成空格
        one = re.sub(r"\s*\n\s*", " ", para).strip()
        doc.add_paragraph(one)

def add_header_block(doc: Document, topic_cn: str, segment_id: int,
                     band: str = DEFAULT_BAND,
                     style: str = DEFAULT_STYLE,
                     role: str = DEFAULT_ROLE):

    # ✅ Header 区（Key 加粗）
    _add_kv(doc, "EXAM", EXAM_NAME)
    _add_kv(doc, "BAND", band)
    _add_kv(doc, "ROLE", role)
    _add_kv(doc, "STYLE", style)
    _add_kv(doc, "TOPIC_CN", topic_cn)
    _add_kv(doc, "TOPIC_EN", "")  # 先留空
    _add_kv(doc, "SEGMENT", f"{segment_id:02d}")

    _add_divider(doc)

def maybe_add_header_block(
    doc: Document,
    topic_cn: str,
    segment_id: int,
    band: str,
    style: str,
    role: str,
):
    """
    Header 总开关封装：
    - SHOW_HEADER = False 时，什么都不写
    - SHOW_HEADER = True 时，才真正调用 add_header_block
    """
    if not SHOW_HEADER:
        return
    add_header_block(
        doc,
        topic_cn=topic_cn,
        segment_id=segment_id,
        band=band,
        style=style,
        role=role,
    )

def maybe_add_header_block_A(
    doc: Document,
    topic_cn: str,
    segment_id: int,
    band: str,
    style: str,
    role: str,
    options: Optional[ExportOptions] = None,
):
    # options 优先生效；否则退回全局开关
    if options is not None:
        if not options.show_header:
            return
    else:
        if not SHOW_HEADER_A:
            return
    add_header_block(doc, topic_cn=topic_cn, segment_id=segment_id, band=band, style=style, role=role)


def maybe_add_header_block_B(
    doc: Document,
    topic_cn: str,
    segment_id: int,
    band: str,
    style: str,
    role: str,
    options: Optional[ExportOptions] = None,
):
    if options is not None:
        if not options.show_header:
            return
    else:
        if not SHOW_HEADER_B:
            return
    add_header_block(doc, topic_cn=topic_cn, segment_id=segment_id, band=band, style=style, role=role)


def maybe_add_header_block_C(
    doc: Document,
    topic_cn: str,
    segment_id: int,
    band: str,
    style: str,
    role: str,
    options: Optional[ExportOptions] = None,
):
    if options is not None:
        if not options.show_header:
            return
    else:
        if not SHOW_HEADER_C:
            return
    add_header_block(doc, topic_cn=topic_cn, segment_id=segment_id, band=band, style=style, role=role)

def _qa_get(qa, key: str, default=""):
    """兼容 qa 是对象 / dataclass / dict"""
    if qa is None:
        return default
    if isinstance(qa, dict):
        v = qa.get(key, default)
    else:
        v = getattr(qa, key, default)
    if v is None:
        return default
    return str(v).strip()

def _qa_question(qa) -> str:
    # 兼容多种字段名
    return (
        _qa_get(qa, "q")
        or _qa_get(qa, "question")
        or _qa_get(qa, "prompt")
        or _qa_get(qa, "Q")
        or _qa_get(qa, "Question")
    )

def add_part1(doc: Document, part1: List[QAItem], options: Optional[ExportOptions] = None):
    opts = options or ExportOptions()
    doc.add_heading("PART 1", level=2)
    _add_divider(doc)

    for idx, qa in enumerate(part1, start=1):
        # Q
        p = doc.add_paragraph()
        question = _qa_question(qa)
        run = p.add_run(f"Q{idx}. {question}")
        run.bold = True

        # Outline（可关）
        if opts.show_outline and qa.outline:
            doc.add_paragraph(f"Outline: {qa.outline}")

        # Answer（可关）
        if opts.show_answers and qa.answer:
            doc.add_paragraph(f"Sample Answer: {qa.answer}")

        doc.add_paragraph("")


def add_part2(doc: Document, part2: Part2Block, options: Optional[ExportOptions] = None):
    opts = options or ExportOptions()
    doc.add_heading("PART 2", level=2)
    _add_divider(doc)

    # ===== Part 2 Cue Card（必须显示） =====
    cue = (getattr(part2, "cue", "") or "").strip()
    bullets = getattr(part2, "bullets", []) or []

    if cue:
        p = doc.add_paragraph("Cue Card:")
        p.runs[0].bold = True

        doc.add_paragraph(cue)
        doc.add_paragraph("")  # 空行，增强结构感

    if bullets:
        p = doc.add_paragraph("You should say:")
        p.runs[0].bold = True
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph("")

    # Outline（可关）
    if opts.show_outline and getattr(part2, "outline", ""):
        doc.add_paragraph(f"Outline: {part2.outline}")

    # Answer（可关）
    if opts.show_answers and getattr(part2, "answer", ""):
        doc.add_paragraph(f"Sample Answer: {part2.answer}")

    doc.add_paragraph("")


def add_part3(doc: Document, part3: List[QAItem], options: Optional[ExportOptions] = None):
    opts = options or ExportOptions()
    doc.add_heading("PART 3", level=2)
    _add_divider(doc)

    for idx, qa in enumerate(part3, start=1):
        p = doc.add_paragraph()
        question = _qa_question(qa)
        run = p.add_run(f"Q{idx}. {question}")
        run.bold = True

        if opts.show_outline and qa.outline:
            doc.add_paragraph(f"Outline: {qa.outline}")

        if opts.show_answers and qa.answer:
            doc.add_paragraph(f"Sample Answer: {qa.answer}")

        doc.add_paragraph("")

def _add_multiline(doc: Document, text: str):
    """
    将多行文本按行写入 docx。
    - 以 '-' 开头的行保留为列表样式（简单用段落前缀）
    - 其它行按普通段落
    """
    text = (text or "").strip()
    if not text:
        doc.add_paragraph("- (empty)")
        return

    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip() != ""]
    for ln in lines:
        doc.add_paragraph(ln)


def render_segment_docx(
    seg: SegmentData,
    out_path: str,
    band: str = DEFAULT_BAND,
    style: str = DEFAULT_STYLE,
    role: str = DEFAULT_ROLE,
    options: Optional[ExportOptions] = None,
):
    opts = options or ExportOptions()

    doc = Document()
    _set_doc_style(doc)

    # ✅ 真实标题：Topic（不要删）
    doc.add_heading(f"{seg.topic}", level=1)

    # ✅ Header(KV)（可开关）
    maybe_add_header_block_B(
        doc,
        topic_cn=seg.topic,
        segment_id=seg.segment_id,
        band=band,
        style=style,
        role=role,
        options=opts,
    )

    # Part 1
    if seg.part1:
        add_part1(doc, seg.part1, options=opts)

    # page_break = "part"：Part1 和 Part23 分页
    if opts.page_break == "part" and seg.part1 and (seg.part2 or seg.part3):
        doc.add_page_break()

    # Part 2
    if seg.part2:
        add_part2(doc, seg.part2, options=opts)

    # page_break = "part"：Part2 和 Part3 分页
    if opts.page_break == "part" and seg.part2 and seg.part3:
        doc.add_page_break()

    # Part 3
    if seg.part3:
        add_part3(doc, seg.part3, options=opts)

    ensure_dir(os.path.dirname(out_path))
    doc.save(out_path)

def add_segment_title(doc: Document, seg, prefix_with_id: bool = True, level: int = 2):
    """seg: SegmentDoc；在合并讲义中打印每个 topic 的小标题"""
    if prefix_with_id:
        title = f"S{seg.segment_id:02d} {seg.topic}"
    else:
        title = f"{seg.topic}"
    doc.add_heading(title, level=level)

def render_merged_docx(segments: list, out_path: str,
                       band: str = DEFAULT_BAND,
                       style: str = DEFAULT_STYLE,
                       role: str = DEFAULT_ROLE,
                       options: Optional[ExportOptions] = None):
    """
    Version B 总讲义（merge_all=True）：
    - 自动分区：先 Part 1，再 Part 2&3
    - 每个 segment 的排版沿用 add_part1/add_part2/add_part3
    - 分页由 options.page_break 控制
    """
    doc = Document()
    _set_doc_style(doc)
    opts = options or ExportOptions()

    # 1) 分组：先 P1 后 P23
    p1_segs = []
    p23_segs = []

    for seg in segments:
        has_p1 = bool(getattr(seg, "part1", None))
        has_p23 = bool(getattr(seg, "part2", None)) or bool(getattr(seg, "part3", None))

        if has_p1 and not has_p23:
            p1_segs.append(seg)
        elif has_p23 and not has_p1:
            p23_segs.append(seg)
        else:
            if has_p1:
                p1_segs.append(seg)
            if has_p23:
                p23_segs.append(seg)

    p1_segs.sort(key=lambda s: getattr(s, "segment_id", 0))
    p23_segs.sort(key=lambda s: getattr(s, "segment_id", 0))

    # 2) 总标题页
    doc.add_heading(f"{EXAM_NAME} Speaking — Version B", level=1)
    _add_kv(doc, "BAND", band)
    _add_kv(doc, "ROLE", role)
    _add_kv(doc, "STYLE", style)
    _add_divider(doc)

    # 3) Part 1 区块
    if p1_segs:
        doc.add_heading("PART 1 — All Topics", level=1)
        _add_divider(doc)

        for i, seg in enumerate(p1_segs):
            doc.add_heading(f"{seg.topic}", level=2)

            maybe_add_header_block_B(
                doc,
                topic_cn=seg.topic,
                segment_id=seg.segment_id,
                band=band,
                style=style,
                role=role,
                options=opts,
            )

            add_part1(doc, seg.part1 or [], options=opts)

            if i != len(p1_segs) - 1:
                if opts.page_break in ("segment", "part"):
                    doc.add_page_break()
                else:
                    _add_divider(doc)

    # 两大区块之间
    if p1_segs and p23_segs:
        if opts.page_break in ("segment", "part"):
            doc.add_page_break()
        else:
            _add_divider(doc)

    # 4) Part 2&3 区块
    if p23_segs:
        doc.add_heading("PART 2 & PART 3 — All Topics", level=1)
        _add_divider(doc)

        for i, seg in enumerate(p23_segs):
            doc.add_heading(f"{seg.topic}", level=2)

            maybe_add_header_block_B(
                doc,
                topic_cn=seg.topic,
                segment_id=seg.segment_id,
                band=band,
                style=style,
                role=role,
                options=opts,
            )

            if getattr(seg, "part2", None):
                add_part2(doc, seg.part2, options=opts)
            if getattr(seg, "part3", None):
                add_part3(doc, seg.part3 or [], options=opts)

            if i != len(p23_segs) - 1:
                if opts.page_break in ("segment", "part"):
                    doc.add_page_break()
                else:
                    _add_divider(doc)
    # 最后整理布局
    _postprocess_docx_layout(doc, keep_page_breaks=False)
    # 保存
    doc.save(out_path)

# -------------------------
# 对外主函数（给 API 调用）
# -------------------------

def export_version_b_docx(run_dir: str,
                          band: str = DEFAULT_BAND,
                          role: str = DEFAULT_ROLE,
                          style: str = DEFAULT_STYLE,
                          merge_all: bool = True,
                          selected_segments: Optional[List[int]] = None,
                          options: Optional[ExportOptions] = None) -> Tuple[str, List[str]]:
    """
    run_dir: runs/<run_id> 的绝对路径
    返回：export_dir, files（文件名列表）
    """
    answers_dir = os.path.join(run_dir, "answers")
    opts = options or ExportOptions()
    export_dir = ensure_dir(os.path.join(run_dir, "exports"))

    seg_map = load_segments_from_answers_dir(answers_dir)
    seg_ids = sorted(seg_map.keys())

    if selected_segments:
        seg_ids = [sid for sid in seg_ids if sid in set(selected_segments)]

    segments = [seg_map[sid] for sid in seg_ids]

    files: List[str] = []

    # 默认：每 segment 一个 docx
    for seg in segments:
        fname = f"S{seg.segment_id:02d}_{safe_filename(seg.topic)}_IELTS_SPK_{band}.docx"
        out_path = os.path.join(export_dir, fname)
        render_segment_docx(seg, out_path, band=band, style=style, role=role, options=opts)
        files.append(fname)

    # merge_all：额外生成一个总讲义
    if merge_all and segments:
        merged_name = f"IELTS_SPK_{band}_All_Topics.docx"
        merged_path = os.path.join(export_dir, merged_name)
        render_merged_docx(segments, merged_path, band=band, style=style, role=role, options=opts)
        files.append(merged_name)

    return export_dir, files

# ============================================================
# Version A: 纯题目版（Questions Only）
# - 每 segment 一个 docx（默认）
# - merge_all=True 额外输出总讲义
# - 数据来源：runs/<run_id>/answers/ 下的 Sxx_P1_*.txt 与 Sxx_P23_*.txt
#   （沿用你现有 parse_part1_file / parse_part23_file / load_segments_from_answers_dir）
# ============================================================

def add_part1_questions_only(
    doc: Document,
    part1: List[QAItem],
    options: Optional["ExportOptions"] = None,
):
    opts = options or ExportOptions()

    """
    Version A - Part 1：只输出题目
    """
    doc.add_heading("PART 1", level=2)
    _add_divider(doc)

    for i, item in enumerate(part1, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"Q{i}. {item.question}")
        r.bold = True

    doc.add_paragraph("")


def add_part2_questions_only(
    doc: Document,
    part2: Part2Block,
    options: Optional["ExportOptions"] = None,
):
    opts = options or ExportOptions()

    """
    Version A - Part 2：输出 Cue Card + bullets（You should say）
    """
    doc.add_heading("PART 2", level=2)
    _add_divider(doc)

    p = doc.add_paragraph()
    p.add_run("Cue Card").bold = True
    doc.add_paragraph((part2.cue or "").strip())

    if part2.bullets:
        p2 = doc.add_paragraph()
        p2.add_run("You should say").bold = True
        _add_bullets(doc, part2.bullets)

    doc.add_paragraph("")


def add_part3_questions_only(
    doc: Document,
    part3: List[QAItem],
    options: Optional["ExportOptions"] = None,
):
    opts = options or ExportOptions()

    """
    Version A - Part 3：只输出题目
    """
    doc.add_heading("PART 3", level=2)
    _add_divider(doc)

    for i, item in enumerate(part3, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"Q{i}. {item.question}")
        r.bold = True

    doc.add_paragraph("")


def render_segment_docx_vA(
    seg: SegmentDoc,
    out_path: str,
    band: str = DEFAULT_BAND,
    style: str = DEFAULT_STYLE,
    role: str = DEFAULT_ROLE,
    options: Optional[ExportOptions] = None):
    """
    Version A - 单个 segment 输出（纯题目版）
    """
    doc = Document()
    _set_doc_style(doc)
    opts = options or ExportOptions()

    # 标题（保持与你 Version B 一致：H1 用 topic）
    doc.add_heading(f"{seg.topic}", level=1)

    # Header 区（字段名固定，ROLE 已是 High School Student (Hangzhou, China)）
    maybe_add_header_block_A(
        doc,
        topic_cn=seg.topic,
        segment_id=seg.segment_id,
        band=band,
        style=style,
        role=role,
    )

    # 仅题目
    if seg.part1:
        add_part1_questions_only(doc, seg.part1, options=opts)

    if seg.part2:
        add_part2_questions_only(doc, seg.part2, options=opts)

    if seg.part3:
        add_part3_questions_only(doc, seg.part3, options=opts)

    ensure_dir(os.path.dirname(out_path))

    # 保存
    doc.save(out_path)


def render_merged_docx_vA(
    segments: List[SegmentDoc],
    out_path: str,
    band: str = DEFAULT_BAND,
    style: str = DEFAULT_STYLE,
    role: str = DEFAULT_ROLE,
    options: Optional[ExportOptions] = None
):
    """
    Version A - 总讲义（纯题目版）
    规则：先 Part 1 topics，再 Part 2&3 topics（与 Version B 的“分区逻辑”一致）
    """
    doc = Document()
    _set_doc_style(doc)
    opts = options or ExportOptions()

    # 封面/总信息
    doc.add_heading(f"{EXAM_NAME} Speaking — Version A (Questions Only)", level=1)
    _add_kv(doc, "BAND", band)
    _add_kv(doc, "ROLE", role)
    _add_kv(doc, "STYLE", style)
    _add_divider(doc)

    # 分组：P1-only / P23-only（防止混杂）
    p1_segs = []
    p23_segs = []
    for seg in segments:
        has_p1 = bool(getattr(seg, "part1", None))
        has_p23 = bool(getattr(seg, "part2", None)) or bool(getattr(seg, "part3", None))
        if has_p1:
            p1_segs.append(seg)
        if has_p23:
            p23_segs.append(seg)

    p1_segs.sort(key=lambda s: getattr(s, "segment_id", 0))
    p23_segs.sort(key=lambda s: getattr(s, "segment_id", 0))

    # PART 1 区块
    if p1_segs:
        doc.add_heading("PART 1 — All Topics", level=1)
        _add_divider(doc)

        for idx, seg in enumerate(p1_segs, start=1):
            doc.add_heading(f"{seg.topic}", level=2)
            maybe_add_header_block_A(
                doc,
                topic_cn=seg.topic,
                segment_id=seg.segment_id,
                band=band,
                style=style,
                role=role,
            )
            add_part1_questions_only(doc, seg.part1 or [], options=opts)
            maybe_page_break(doc, opts)

    # PART 2&3 区块
    if p23_segs:
        doc.add_heading("PART 2 & PART 3 — All Topics", level=1)
        _add_divider(doc)

        for idx, seg in enumerate(p23_segs, start=1):
            doc.add_heading(f"{seg.topic}", level=2)
            maybe_add_header_block_A(
                doc,
                topic_cn=seg.topic,
                segment_id=seg.segment_id,
                band=band,
                style=style,
                role=role,
            )
            if seg.part2:
                add_part2_questions_only(doc, seg.part2, options=opts)
            if seg.part3:
                add_part3_questions_only(doc, seg.part3, options=opts)
            maybe_page_break(doc, opts)

    ensure_dir(os.path.dirname(out_path))

    # 最后整理布局
    _postprocess_docx_layout(doc, keep_page_breaks=False)
    # 保存
    doc.save(out_path)


def export_version_a_docx(
    run_dir: str,
    band: str = DEFAULT_BAND,
    role: str = DEFAULT_ROLE,
    style: str = DEFAULT_STYLE,
    merge_all: bool = True,
    selected_segments: Optional[List[int]] = None,
    options: Optional[ExportOptions] = None) -> Tuple[str, List[str]]:
    """
    对外主函数（给 API 调用）—— Version A（纯题目版）
    参数/返回结构刻意对齐 export_version_b_docx，减少联动成本。
    """
    answers_dir = os.path.join(run_dir, "answers")
    opts = options or ExportOptions()
    export_dir = ensure_dir(os.path.join(run_dir, "exports"))

    seg_map = load_segments_from_answers_dir(answers_dir)
    seg_ids = sorted(seg_map.keys())

    if selected_segments:
        wanted = set(int(x) for x in selected_segments)
        seg_ids = [sid for sid in seg_ids if sid in wanted]

    segments = [seg_map[sid] for sid in seg_ids]

    files: List[str] = []

    # 默认：每 segment 一个 docx
    for seg in segments:
        fname = f"S{seg.segment_id:02d}_{safe_filename(seg.topic)}_IELTS_SPK_{band}_V_A.docx"
        out_path = os.path.join(export_dir, fname)
        render_segment_docx_vA(seg, out_path, band=band, style=style, role=role, options=opts)
        files.append(fname)

    # merge_all：额外生成一个总讲义
    if merge_all and segments:
        merged_name = f"IELTS_SPK_{band}_All_Topics_V_A.docx"
        merged_path = os.path.join(export_dir, merged_name)
        render_merged_docx_vA(segments, merged_path, band=band, style=style, role=role, options=opts)
        files.append(merged_name)

    return export_dir, files

# =========================
# Version C（可截图宣传版）
# - 每个 Part 顶部都显示 Topic
# - 只输出 Q + Sample Answer（不输出 Outline）
# - 一个 Part 一页；超页自动拆分并加 (1)(2)...
# =========================

# 你可以按你实际版面微调这三个阈值（越小越容易分成多页，越大越容易挤爆一页）
VC_P1_QA_PER_PAGE = 5   # Part 1 每页最多几个 Q&A
VC_P3_QA_PER_PAGE = 3   # Part 3 每页最多几个 Q&A
VC_P2_WORDS_PER_PAGE = 500  # Part 2 Answer 每页大约多少词（粗略分页，够用）

def _count_words(s: str) -> int:
    return len(re.findall(r"[A-Za-z0-9']+", s or ""))

def _split_list_into_pages(items, n_per_page: int):
    if not items:
        return []
    n_per_page = max(int(n_per_page), 1)
    return [items[i:i+n_per_page] for i in range(0, len(items), n_per_page)]

def _split_text_by_words(text: str, words_per_page: int) -> List[str]:
    """
    把一大段 answer 粗略按“词数”切成多页。
    规则：优先按句号/问号/叹号切句，保持自然；实在不行才硬切。
    """
    text = (text or "").strip()
    if not text:
        return [""]

    words_per_page = max(int(words_per_page), 80)

    # 先按换行段落拆（如果模型输出有分段，会更稳定）
    paras = [p.strip() for p in text.splitlines() if p.strip()]
    if len(paras) >= 2:
        # 逐段累积到阈值
        pages, cur, cur_w = [], [], 0
        for p in paras:
            w = _count_words(p)
            if cur and (cur_w + w > words_per_page):
                pages.append("\n".join(cur).strip())
                cur, cur_w = [], 0
            cur.append(p)
            cur_w += w
        if cur:
            pages.append("\n".join(cur).strip())
        return pages

    # 没有自然段：按句子拆
    sentences = re.split(r"(?<=[.!?])\s+", text)
    pages, cur, cur_w = [], [], 0
    for sent in sentences:
        sent = sent.strip()
        if not sent:
            continue
        w = _count_words(sent)
        if cur and (cur_w + w > words_per_page):
            pages.append(" ".join(cur).strip())
            cur, cur_w = [], 0
        cur.append(sent)
        cur_w += w

    if cur:
        pages.append(" ".join(cur).strip())

    # 兜底：如果被切成空，返回原文
    return pages or [text]


def _add_part_heading_vc(doc: Document, part_label: str, topic: str, page_idx: int, total_pages: int):
    """
    part_label: "PART 1" / "PART 2" / "PART 3"
    """
    suffix = f" ({page_idx})" if total_pages > 1 else ""
    doc.add_heading(f"{part_label} — {topic}{suffix}", level=1)


def _add_q_and_answer_vc(doc: Document, q: str, answer: str,
                         outline: str = "",
                         options: Optional[ExportOptions] = None):
    opts = options or ExportOptions()

    # Q
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True

    # Answer（可选）
    if opts.show_answers:
        ap = doc.add_paragraph(answer or "")
        ap.paragraph_format.space_after = Pt(6)

def add_part1_vC(
    doc: Document,
    topic: str,
    items: List[QAItem],
    options: Optional[ExportOptions] = None
):
    opts = options or ExportOptions()

    pages = _split_list_into_pages(items, VC_P1_QA_PER_PAGE)

    for pi, page_items in enumerate(pages, start=1):
        _add_part_heading_vc(doc, "PART 1", topic, pi, len(pages))

        # ✅ 全局序号 offset：保证第 2 页从 Q(VC_P1_QA_PER_PAGE+1) 开始
        base = (pi - 1) * VC_P1_QA_PER_PAGE

        for local_idx, item in enumerate(page_items, start=1):
            q_idx = base + local_idx
            _add_q_and_answer_vc(
                doc,
                f"Q{q_idx}. {getattr(item, 'question', '')}",
                getattr(item, "answer", ""),
                outline=getattr(item, "outline", ""),
                options=opts,
            )

        # ✅ 只在“不是最后一页”时分页；并且尊重 page_break 开关
        if opts.page_break != "none" and pi < len(pages):
            doc.add_page_break()

def add_part3_vC(
    doc: Document,
    topic: str,
    items: List[QAItem],
    options: Optional[ExportOptions] = None,
):
    opts = options or ExportOptions()
    pages = _split_list_into_pages(items, VC_P3_QA_PER_PAGE)

    for pi, page_items in enumerate(pages, start=1):
        _add_part_heading_vc(doc, "PART 3", topic, pi, len(pages))

        # ✅ 关键：全局序号 offset，保证第 2 页从 Q4 开始（若每页 3 题）
        base = (pi - 1) * VC_P3_QA_PER_PAGE

        for local_idx, item in enumerate(page_items, start=1):
            q_idx = base + local_idx
            _add_q_and_answer_vc(
                doc,
                f"Q{q_idx}. {getattr(item, 'question', '')}",
                getattr(item, "answer", ""),
                outline=getattr(item, "outline", ""),
                options=opts,
            )

        if opts.page_break != "none" and pi < len(pages):
            doc.add_page_break()

def add_part2_vC(
    doc: Document,
    topic: str,
    part2: Part2Block,
    options: Optional[ExportOptions] = None,
):
    opts = options or ExportOptions()

    cue = (getattr(part2, "cue", "") or "").strip()
    bullets = getattr(part2, "bullets", None) or []
    answer_text = getattr(part2, "answer", "") or ""
    outline_text = getattr(part2, "outline", "") or ""

    answer_pages = _split_text_by_words(answer_text, VC_P2_WORDS_PER_PAGE)

    # 防御：空答案也至少出 1 页，避免 len(answer_pages)=0 导致标题不显示
    if not answer_pages:
        answer_pages = [""]

    for pi, ans_chunk in enumerate(answer_pages, start=1):
        _add_part_heading_vc(doc, "PART 2", topic, pi, len(answer_pages))

        # cue + bullets：每一页都显示（方便截图）
        if cue:
            p = doc.add_paragraph("Cue Card:")
            p.runs[0].bold = True
            doc.add_paragraph(cue)

        if bullets:
            p = doc.add_paragraph("You should say:")
            p.runs[0].bold = True
            for b in bullets:
                doc.add_paragraph(f"- {b}")

        # ✅ V-C：不需要独立 Outline 区，但你要“开关”，所以这里交给 _add_q_and_answer_vc
        # 用一个“虚拟问题标题”来复用同一套渲染逻辑（最小耦合）
        _add_q_and_answer_vc(
            doc,
            "Sample Answer",
            ans_chunk,
            outline=outline_text,
            options=opts,
        )

        if opts.page_break != "none" and pi < len(answer_pages):
            doc.add_page_break()

def render_segment_docx_vC(seg: SegmentDoc,
                           out_path: str,
                           band: str = DEFAULT_BAND,
                           style: str = DEFAULT_STYLE,
                           role: str = DEFAULT_ROLE,
                           options: Optional[ExportOptions] = None):
    """
    Version C（截图/练习用）：
    - 每个 Part 标题都带 Topic：例如 “PART 1 — Advertisement (1)”
    - 默认 Q + Sample Answer（可用 options 开关控制）
    - page_break 模式：
        - none：不分页
        - part：Part 1 / Part 2 / Part 3 各自分页（并支持 (1)(2)...）
        - segment：对单文件来说等价于 part（保守处理）
    """
    doc = Document()
    _set_doc_style(doc)
    opts = options or ExportOptions()

    topic = getattr(seg, "topic", "") or "UnknownTopic"

    wrote_any = False
    if seg.part1:
        add_part1_vC(doc, topic, seg.part1, options=opts)
        wrote_any = True

    if seg.part2:
        if wrote_any and opts.page_break in ("part", "segment"):
            doc.add_page_break()
        add_part2_vC(doc, topic, seg.part2, options=opts)
        wrote_any = True

    if seg.part3:
        if wrote_any and opts.page_break in ("part", "segment"):
            doc.add_page_break()
        add_part3_vC(doc, topic, seg.part3, options=opts)
        wrote_any = True

    ensure_dir(os.path.dirname(out_path))
    _postprocess_docx_c_layout(doc, add_page_number=False)  # V-C 截图版：默认不加页码
    doc.save(out_path)

def render_merged_docx_vC(segments: List[SegmentDoc],
                          out_path: str,
                          band: str = DEFAULT_BAND,
                          style: str = DEFAULT_STYLE,
                          role: str = DEFAULT_ROLE,
                          options: Optional[ExportOptions] = None):
    """
    Version C merge_all：把所有 segment 合并到一个 docx
    - segment 之间是否分页：由 options.page_break 控制
    - segment 内：Part 1 / Part 2 / Part 3 分页同 add_part*_vC 的逻辑
    """
    doc = Document()
    _set_doc_style(doc)
    opts = options or ExportOptions()

    for si, seg in enumerate(segments, start=1):
        if si > 1 and opts.page_break in ("segment", "part"):
            doc.add_page_break()

        topic = getattr(seg, "topic", "") or "UnknownTopic"

        wrote_any = False
        if seg.part1:
            add_part1_vC(doc, topic, seg.part1, options=opts)
            wrote_any = True

        if seg.part2:
            if wrote_any and opts.page_break in ("part", "segment"):
                doc.add_page_break()
            add_part2_vC(doc, topic, seg.part2, options=opts)
            wrote_any = True

        if seg.part3:
            if wrote_any and opts.page_break in ("part", "segment"):
                doc.add_page_break()
            add_part3_vC(doc, topic, seg.part3, options=opts)
            wrote_any = True

    ensure_dir(os.path.dirname(out_path))
    _postprocess_docx_c_layout(doc, add_page_number=False)

    # 保存
    doc.save(out_path)

def export_version_c_docx(run_dir: str,
                          band: str = DEFAULT_BAND,
                          role: str = DEFAULT_ROLE,
                          style: str = DEFAULT_STYLE,
                          merge_all: bool = True,
                          selected_segments: Optional[List[int]] = None,
                          options: Optional[ExportOptions] = None) -> Tuple[str, List[str]]:
    """
    Version C 导出：
    - 每 segment 一个 docx
    - merge_all=True 额外生成一个合并 docx
    """
    answers_dir = os.path.join(run_dir, "answers")
    opts = options or ExportOptions()
    export_dir = os.path.join(run_dir, "exports")
    ensure_dir(export_dir)

    seg_map = load_segments_from_answers_dir(answers_dir)  # {segment_id: SegmentDoc}
    seg_ids = sorted(seg_map.keys())

    if selected_segments:
        selected_set = set(int(x) for x in selected_segments)
        seg_ids = [sid for sid in seg_ids if sid in selected_set]

    segments = [seg_map[sid] for sid in seg_ids]

    files: List[str] = []
    for seg in segments:
        fname = f"S{seg.segment_id:02d}_{_safe_filename(seg.topic)}_V_C.docx"
        out_path = os.path.join(export_dir, fname)
        render_segment_docx_vC(seg, out_path, band=band, style=style, role=role, options=opts)
        files.append(fname)

    if merge_all and segments:
        merged_name = f"IELTS_SPK_{band}_All_Topics_V_C.docx"
        merged_path = os.path.join(export_dir, merged_name)
        render_merged_docx_vC(segments, merged_path, band=band, style=style, role=role, options=opts)
        files.append(merged_name)

    return export_dir, files


# 排版优化辅助函数
PART_TITLE_RE = re.compile(r"^\s*PART\s*(1|2|3)\s*$", re.I)

# 排版1：安全删除段落
def _remove_paragraph(p):
    """安全删除段落（只删这个段落，不动其它正文）"""
    el = p._element
    el.getparent().remove(el)
# 排版1.1：在段落后插入新段落
def _insert_paragraph_after(paragraph, text: str = "", style=None):
    """
    在指定 paragraph 后插入一个新段落，并返回该新段落（python-docx 原生缺这个函数）。
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    # 上面 add_paragraph() 会追加到末尾，所以我们需要“绑定”到 new_p
    new_para._p = new_p  # 关键：让 Paragraph 对象指向我们插入的位置

    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

# 排版2：段落迭代器
def _iter_paragraphs(doc):
    # python-docx 的 doc.paragraphs 会动态变化；这里每次取最新引用更稳
    return list(doc.paragraphs)
#排版3：判断空段落
def _is_empty_para(p) -> bool:
    # 只删“真正的空段落”：无文本、无可见 run 内容
    return not (p.text or "").strip()
# 排版4：设置字体
def _set_run_font(run, name="等线", size=11):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
# 排版5：设置段落字体
def _set_paragraph_font(p, name="等线", size=11):
    for r in p.runs:
        _set_run_font(r, name=name, size=size)
# 排版6：页脚页码
def _ensure_page_number_centered(doc):
    """页脚居中页码：保留（不依赖模板）"""
    section = doc.sections[0]
    footer = section.footer

    # 用第一个段落；如果没有则创建
    if footer.paragraphs:
        fp = footer.paragraphs[0]
        # 清空现有 runs（只影响页脚，不碰正文）
        for r in list(fp.runs):
            r._element.getparent().remove(r._element)
    else:
        fp = footer.add_paragraph()

    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    _set_run_font(run)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
# 排版7：移除所有分页符
def _remove_all_page_breaks(doc):
    """
    移除正文中的分页符（V-A/B 要求不分页）。
    处理：run 里的 <w:br w:type="page"/> + 段落属性中的 pageBreakBefore。
    """
    for p in _iter_paragraphs(doc):
        # 段落级：pageBreakBefore
        pPr = p._p.pPr
        if pPr is not None:
            pbb = pPr.find(qn("w:pageBreakBefore"))
            if pbb is not None:
                pPr.remove(pbb)

        # run级：<w:br w:type="page"/>
        for r in p.runs:
            brs = r._r.findall(qn("w:br"))
            for br in brs:
                if br.get(qn("w:type")) == "page":
                    r._r.remove(br)
# 排版8：处理 Outline 段落
def _split_outline_para_to_title_and_bullets(doc, p):
    """
    把：
      Outline: - A
      - B
      - C
    合并成同一个段落：
      Outline:
      - A
      - B
      - C

    不创建 bullet 列表，不插入段落，避免一切结构风险。
    """
    txt0 = (p.text or "").strip()
    if not txt0.lower().startswith("outline:"):
        return

    # 1) 取出 Outline: 后面同一行的第一条（如果存在）
    tail = txt0[len("Outline:"):].replace("\t", " ").strip()

    def norm_dash(s: str) -> str:
        s = (s or "").replace("\t", " ").strip()
        s = s.lstrip("•").strip()
        if s.startswith("-"):
            s = s.lstrip("-").strip()
        return f"- {s}" if s else ""

    lines = ["Outline:"]
    first = norm_dash(tail)
    if first:
        lines.append(first)

    # 2) 把紧随其后的连续 "- xxx" 段落收集进来，然后删掉那些段落
    paras = _iter_paragraphs(doc)
    try:
        i = paras.index(p)
    except ValueError:
        return

    j = i + 1
    to_delete = []
    while j < len(paras):
        q = paras[j]
        t = (q.text or "").strip()
        if not t:
            # 空行先停（避免吞掉不该吞的段落）
            break
        if t.startswith("-") or t.startswith("•"):
            lines.append(norm_dash(t))
            to_delete.append(q)
            j += 1
            continue
        break

    # 3) 写回同一个段落（换行显示）
    p.text = "\n".join(lines)
    _set_paragraph_font(p)

    # 4) 删除已合并的后续 "- xxx" 段落
    for q in to_delete:
        _remove_paragraph(q)


# 排版9：处理 Sample Answer 段落
def _split_sample_answer_to_title_and_body(doc, p):
    """
    输入：
      Sample Answer: xxx...
    输出：
      Sample Answer:
      xxx...
    注意：不拆分正文段落，不合并，不删内容。
    """
    txt = (p.text or "").strip()
    m = re.match(r"^Sample Answer:\s*(.*)$", txt, flags=re.I)
    if not m:
        return

    body = (m.group(1) or "").strip()
    p.text = "Sample Answer:"
    _set_paragraph_font(p)

    if body:
        new_p = _insert_paragraph_after(p, body, style=doc.styles["Normal"])
        _set_paragraph_font(new_p)
# 排版主函数,针对V-AB
def _postprocess_docx_layout(doc, *, keep_page_breaks: bool):
    """
    只做你要求的：
    - 删除：PART 1/2/3（所有 H2 标题那行直接去掉）
    - 删除：多余空行（^p^p 等价于连续空段落）
    - 处理：Outline / Sample Answer 版式
    - 字体：等线 11
    - 页码：页脚居中保留
    - 分页：V-A/B 全部去掉分页；V-C 保留分页
    关键：绝不“扫描式清空正文”，只删两类段落（空段落、PART标题段落）。
    """

    # 1) 可选：去分页（只对 V-A/B）
    if not keep_page_breaks:
        _remove_all_page_breaks(doc)

    # 2) 遍历并处理（用 while + 每轮重取列表，避免 doc.paragraphs 变化导致错位）
    i = 0
    while True:
        paras = _iter_paragraphs(doc)
        if i >= len(paras):
            break

        p = paras[i]
        txt = (p.text or "").strip()

        # 删除 <<<END>>>（整段删除，不留空行）
        if txt == "<<<END>>>":
            _remove_paragraph(p)
            continue  # 不递增 i，当前位置已变成下一段

        # 2.1 删除空段落（压缩 ^p^p）
        if _is_empty_para(p):
            _remove_paragraph(p)
            continue  # 不递增 i，当前位置已变成下一段

        # 2.2 删除 PART 标题行（你说“所有 H2 标题的 PART 1/2/3 行去掉”）
        #     这里不强依赖样式名，只要文本是 PART 1/2/3 就删，最稳。
        if PART_TITLE_RE.match(txt):
            _remove_paragraph(p)
            continue

        # 2.3 Outline / Sample Answer 排版处理
        if txt.lower().startswith("outline:"):
            _split_outline_para_to_title_and_bullets(doc, p)
            # 处理后继续往后走
        elif txt.lower().startswith("sample answer:"):
            _split_sample_answer_to_title_and_body(doc, p)

        # 2.4 全段字体统一（不动内容）
        _set_paragraph_font(p)

        i += 1

    # 3) 页码（居中保留）
    _ensure_page_number_centered(doc)


#新增排版函数 针对V-C
# =========================
# V-C 专用：更稳的后处理
# - 保留分页符（run 的 page break + 段落级 pageBreakBefore）
# - 不删除 V-C 标题（PART X — topic）
# - 清理真正空段落、<<<END>>>、以及 * / ** 格式标记
# - 全文统一字体：等线 11
# =========================

def _para_has_page_break_run(p) -> bool:
    """run级分页：<w:br w:type="page"/>"""
    for r in p.runs:
        brs = r._r.findall(qn("w:br"))
        for br in brs:
            if br.get(qn("w:type")) == "page":
                return True
    return False

def _para_has_page_break_before(p) -> bool:
    """段落级分页：<w:pageBreakBefore/>"""
    pPr = p._p.pPr
    if pPr is None:
        return False
    return pPr.find(qn("w:pageBreakBefore")) is not None

def _para_has_any_page_break(p) -> bool:
    return _para_has_page_break_run(p) or _para_has_page_break_before(p)

_STAR_EMPH_RE = re.compile(r"(\*\*|\*)([^*\n]+?)\1")  # *word* or **word**
_MULTI_STAR_RE = re.compile(r"\*{2,}")               # ** / *** / ****

def _clean_star_markers(text: str) -> str:
    """
    清理 * / ** 这类 markdown 强调标记：
    - *word*  -> word
    - **word** -> word
    - 残留的 **/*** -> 删除
    注意：这是“导出层”清洗，宁可保守，不做复杂语义判断。
    """
    s = text or ""
    # 先去掉成对的强调
    prev = None
    while prev != s:
        prev = s
        s = _STAR_EMPH_RE.sub(r"\2", s)
    # 再去掉残留的连续星号
    s = _MULTI_STAR_RE.sub("", s)
    # 去掉孤立星号（只在它作为“分隔符/装饰符”时）
    # 规则：两侧是空白或行首/行尾
    s = re.sub(r"(^|\s)\*(\s|$)", r"\1\2", s)
    # 压缩多空白
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s

def _postprocess_docx_c_layout(doc: Document, *, add_page_number: bool = False):
    """
    V-C 专用后处理：
    - 保留分页符（绝不删除）
    - 删除：<<<END>>> 段落、真正空段落（但不删分页段落）
    - 清理：* / ** 格式标记
    - 字体：等线 11
    - 页码：默认不加（截图版一般不要），需要则 add_page_number=True
    """
    i = 0
    while True:
        paras = _iter_paragraphs(doc)
        if i >= len(paras):
            break

        p = paras[i]
        txt = (p.text or "")

        # 1) 删除 <<<END>>>
        if txt.strip() == "<<<END>>>":
            _remove_paragraph(p)
            continue

        # 2) 真正空段落：但要保留分页符段落
        if not txt.strip():
            if _para_has_any_page_break(p):
                # 分页段落不能删
                _set_paragraph_font(p)  # 顺手统一字体（即使不可见）
                i += 1
                continue
            _remove_paragraph(p)
            continue

        # 3) 清理星号格式标记：逐 run 处理更安全（不破坏 runs 结构太多）
        for r in p.runs:
            if r.text:
                r.text = _clean_star_markers(r.text)

        # 4) 段落级再做一次（防止 run 之间拼起来出现 **）
        #    注意：直接写 p.text 会重建 runs（会丢粗体），V-C 一般不依赖 run 级粗体
        cleaned = _clean_star_markers(p.text)
        if cleaned != p.text:
            p.text = cleaned

        # 5) 全段字体统一
        _set_paragraph_font(p, name="等线", size=11)

        i += 1

    if add_page_number:
        _ensure_page_number_centered(doc)
